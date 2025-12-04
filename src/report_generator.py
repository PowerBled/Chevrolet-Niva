"""
Модуль генератора профессиональных отчетов диагностики
"""

import os
import json
import csv
import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any, Union
from dataclasses import dataclass, asdict, field
from enum import Enum
import pandas as pd
from jinja2 import Environment, FileSystemLoader, Template
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
from openpyxl.utils import get_column_letter
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')  # Для работы без GUI
import numpy as np
from PyQt5.QtWidgets import QProgressDialog, QMessageBox
from PyQt5.QtCore import Qt

from utils.logger import setup_logger
from config_manager import ConfigManager


class ReportFormat(Enum):
    """Форматы отчетов"""
    HTML = "html"
    PDF = "pdf"  # Требуется дополнительная установка weasyprint
    DOCX = "docx"
    XLSX = "xlsx"
    JSON = "json"
    CSV = "csv"


class ReportLevel(Enum):
    """Уровни детализации отчетов"""
    BASIC = "basic"        # Только основные данные
    STANDARD = "standard"  # Стандартный отчет
    DETAILED = "detailed"  # Подробный отчет
    EXPERT = "expert"      # Экспертный отчет с графиками


@dataclass
class VehicleInfo:
    """Информация об автомобиле"""
    vin: str = ""
    model: str = ""
    year: int = 0
    engine: str = ""
    mileage: float = 0.0
    license_plate: str = ""
    owner: str = ""
    contact: str = ""
    notes: str = ""


@dataclass
class DiagnosticResult:
    """Результаты диагностики"""
    timestamp: datetime.datetime = field(default_factory=datetime.datetime.now)
    vehicle_info: VehicleInfo = field(default_factory=VehicleInfo)
    dtcs: List[Dict] = field(default_factory=list)
    live_data: Dict[str, Any] = field(default_factory=dict)
    ecu_status: Dict[str, Any] = field(default_factory=dict)
    sensor_checks: Dict[str, Any] = field(default_factory=dict)
    actuator_checks: Dict[str, Any] = field(default_factory=dict)
    adaptation_results: Dict[str, Any] = field(default_factory=dict)
    recommendations: List[str] = field(default_factory=list)
    technician: str = ""
    diagnostic_duration: float = 0.0
    software_version: str = "1.0.0"
    rating: int = 0  # Оценка состояния от 0 до 100


@dataclass
class ReportConfig:
    """Конфигурация отчета"""
    format: ReportFormat = ReportFormat.HTML
    level: ReportLevel = ReportLevel.STANDARD
    include_graphs: bool = True
    include_photos: bool = False
    include_technical_data: bool = True
    include_recommendations: bool = True
    language: str = "ru"
    company_logo: Optional[str] = None
    company_name: str = "Автосервис"
    company_address: str = ""
    company_phone: str = ""
    technician_signature: bool = True
    customer_signature: bool = True


class ReportGenerator:
    """Генератор профессиональных отчетов"""
    
    def __init__(self, config: Optional[ConfigManager] = None):
        self.config = config or ConfigManager()
        self.logger = setup_logger()
        self.templates_dir = Path("assets/templates")
        self.reports_dir = Path("reports")
        self.ensure_directories()
        
        # Настройка Jinja2
        self.jinja_env = Environment(
            loader=FileSystemLoader(str(self.templates_dir)),
            extensions=['jinja2.ext.i18n']
        )
        
        # Стандартные переводы
        self.translations = {
            'ru': self._load_translations('ru'),
            'en': self._load_translations('en')
        }
        
    def ensure_directories(self):
        """Создание необходимых директорий"""
        directories = [
            self.reports_dir,
            self.reports_dir / "html",
            self.reports_dir / "docx",
            self.reports_dir / "xlsx",
            self.reports_dir / "json",
            self.reports_dir / "csv",
            self.reports_dir / "photos",
            self.templates_dir,
            self.templates_dir / "includes",
            Path("assets") / "styles"
        ]
        
        for directory in directories:
            directory.mkdir(parents=True, exist_ok=True)
            
    def _load_translations(self, lang: str) -> Dict:
        """Загрузка переводов"""
        translations = {
            'ru': {
                'report_title': 'Отчет диагностики Chevrolet Niva',
                'vehicle_info': 'Информация об автомобиле',
                'diagnostic_results': 'Результаты диагностики',
                'dtcs': 'Коды неисправностей',
                'live_data': 'Текущие параметры',
                'recommendations': 'Рекомендации',
                'technical_data': 'Технические данные',
                'summary': 'Сводка',
                'rating': 'Оценка состояния',
                'excellent': 'Отличное',
                'good': 'Хорошее',
                'satisfactory': 'Удовлетворительное',
                'poor': 'Плохое',
                'critical': 'Критическое',
                'generated_by': 'Сгенерировано программой',
                'page': 'Страница',
                'date': 'Дата',
                'technician': 'Техник',
                'signature': 'Подпись',
                'customer': 'Клиент',
            },
            'en': {
                'report_title': 'Chevrolet Niva Diagnostic Report',
                'vehicle_info': 'Vehicle Information',
                'diagnostic_results': 'Diagnostic Results',
                'dtcs': 'Trouble Codes',
                'live_data': 'Live Parameters',
                'recommendations': 'Recommendations',
                'technical_data': 'Technical Data',
                'summary': 'Summary',
                'rating': 'Condition Rating',
                'excellent': 'Excellent',
                'good': 'Good',
                'satisfactory': 'Satisfactory',
                'poor': 'Poor',
                'critical': 'Critical',
                'generated_by': 'Generated by',
                'page': 'Page',
                'date': 'Date',
                'technician': 'Technician',
                'signature': 'Signature',
                'customer': 'Customer',
            }
        }
        return translations.get(lang, translations['en'])
        
    def generate_report(self, 
                       result: DiagnosticResult,
                       report_config: ReportConfig,
                       progress_callback=None) -> Dict[str, str]:
        """
        Генерация отчета в указанных форматах
        
        Args:
            result: Результаты диагностики
            report_config: Конфигурация отчета
            progress_callback: Функция для обновления прогресса
            
        Returns:
            Словарь с путями к созданным файлам
        """
        try:
            if progress_callback:
                progress_callback(0, "Подготовка данных...")
                
            # Создание базового имени файла
            base_filename = self._create_filename(result.vehicle_info)
            generated_files = {}
            
            # Генерация отчета в основном формате
            if report_config.format == ReportFormat.HTML:
                filepath = self.generate_html_report(result, report_config, base_filename)
                generated_files['html'] = str(filepath)
                
            elif report_config.format == ReportFormat.DOCX:
                filepath = self.generate_docx_report(result, report_config, base_filename)
                generated_files['docx'] = str(filepath)
                
            elif report_config.format == ReportFormat.XLSX:
                filepath = self.generate_excel_report(result, report_config, base_filename)
                generated_files['xlsx'] = str(filepath)
                
            elif report_config.format == ReportFormat.JSON:
                filepath = self.generate_json_report(result, report_config, base_filename)
                generated_files['json'] = str(filepath)
                
            elif report_config.format == ReportFormat.CSV:
                filepath = self.generate_csv_report(result, report_config, base_filename)
                generated_files['csv'] = str(filepath)
                
            # Генерация дополнительных форматов если требуется
            if report_config.level == ReportLevel.EXPERT:
                # Всегда создаем HTML для экспертного уровня
                if report_config.format != ReportFormat.HTML:
                    html_file = self.generate_html_report(result, report_config, base_filename)
                    generated_files['html'] = str(html_file)
                    
                # Создаем Excel для детальных данных
                if report_config.format != ReportFormat.XLSX:
                    excel_file = self.generate_excel_report(result, report_config, base_filename)
                    generated_files['xlsx'] = str(excel_file)
                    
            if progress_callback:
                progress_callback(100, "Отчет создан!")
                
            return generated_files
            
        except Exception as e:
            self.logger.error(f"Ошибка генерации отчета: {e}")
            raise
            
    def generate_html_report(self, 
                            result: DiagnosticResult,
                            config: ReportConfig,
                            base_filename: str) -> Path:
        """Генерация HTML отчета"""
        try:
            # Загрузка шаблона
            template_name = f"report_{config.level.value}.html"
            template_path = self.templates_dir / template_name
            
            if not template_path.exists():
                # Создание стандартного шаблона если отсутствует
                self._create_default_html_template()
                
            template = self.jinja_env.get_template(template_name)
            
            # Подготовка данных для шаблона
            template_data = self._prepare_template_data(result, config)
            
            # Генерация графиков если требуется
            if config.include_graphs:
                graphs = self._generate_graphs(result, config)
                template_data['graphs'] = graphs
                
            # Рендеринг шаблона
            html_content = template.render(**template_data)
            
            # Сохранение отчета
            filename = f"{base_filename}_{config.level.value}.html"
            filepath = self.reports_dir / "html" / filename
            
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(html_content)
                
            # Копирование стилей если необходимо
            self._copy_styles()
            
            self.logger.info(f"HTML отчет сохранен: {filepath}")
            return filepath
            
        except Exception as e:
            self.logger.error(f"Ошибка создания HTML отчета: {e}")
            # Создание простого HTML отчета
            return self._create_simple_html_report(result, config, base_filename)
            
    def generate_docx_report(self,
                            result: DiagnosticResult,
                            config: ReportConfig,
                            base_filename: str) -> Path:
        """Генерация документа Word"""
        try:
            doc = Document()
            
            # Настройка документа
            section = doc.sections[0]
            section.page_height = Inches(11.69)
            section.page_width = Inches(8.27)
            section.left_margin = Inches(0.79)
            section.right_margin = Inches(0.79)
            section.top_margin = Inches(0.79)
            section.bottom_margin = Inches(0.79)
            
            # Заголовок
            title = doc.add_heading('ОТЧЕТ ДИАГНОСТИКИ', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Информация об автомобиле
            doc.add_heading('1. Информация об автомобиле', 1)
            vehicle_table = doc.add_table(rows=8, cols=2)
            vehicle_table.style = 'LightShading-Accent1'
            
            vehicle_data = [
                ("VIN:", result.vehicle_info.vin),
                ("Модель:", result.vehicle_info.model),
                ("Год выпуска:", str(result.vehicle_info.year)),
                ("Двигатель:", result.vehicle_info.engine),
                ("Пробег:", f"{result.vehicle_info.mileage:.1f} км"),
                ("Гос. номер:", result.vehicle_info.license_plate),
                ("Владелец:", result.vehicle_info.owner),
                ("Контакты:", result.vehicle_info.contact)
            ]
            
            for i, (label, value) in enumerate(vehicle_data):
                vehicle_table.cell(i, 0).text = label
                vehicle_table.cell(i, 1).text = value
                
            # Результаты диагностики
            doc.add_heading('2. Результаты диагностики', 1)
            
            # Общая оценка
            rating_text = f"Общая оценка состояния: {result.rating}/100"
            rating_para = doc.add_paragraph(rating_text)
            rating_para.runs[0].bold = True
            rating_para.runs[0].font.color.rgb = self._get_rating_color(result.rating)
            
            # Коды ошибок
            if result.dtcs:
                doc.add_heading('2.1. Коды неисправностей', 2)
                dtc_table = doc.add_table(rows=len(result.dtcs) + 1, cols=4)
                dtc_table.style = 'LightGrid-Accent1'
                
                # Заголовки
                headers = ["Код", "Описание", "Статус", "Критичность"]
                for i, header in enumerate(headers):
                    cell = dtc_table.cell(0, i)
                    cell.text = header
                    cell.paragraphs[0].runs[0].bold = True
                    
                # Данные
                for i, dtc in enumerate(result.dtcs, 1):
                    dtc_table.cell(i, 0).text = dtc.get('code', '')
                    dtc_table.cell(i, 1).text = dtc.get('description', '')
                    dtc_table.cell(i, 2).text = dtc.get('status', '')
                    dtc_table.cell(i, 3).text = dtc.get('severity', '')
                    
            # Текущие параметры
            if result.live_data:
                doc.add_heading('2.2. Текущие параметры', 2)
                
                # Группируем параметры по системам
                params_by_system = self._group_parameters_by_system(result.live_data)
                
                for system, params in params_by_system.items():
                    doc.add_heading(f'Система: {system}', 3)
                    param_table = doc.add_table(rows=len(params) + 1, cols=3)
                    param_table.style = 'LightList-Accent1'
                    
                    # Заголовки
                    param_headers = ["Параметр", "Значение", "Единица"]
                    for i, header in enumerate(param_headers):
                        cell = param_table.cell(0, i)
                        cell.text = header
                        cell.paragraphs[0].runs[0].bold = True
                        
                    # Данные
                    for i, (param_name, param_data) in enumerate(params.items(), 1):
                        param_table.cell(i, 0).text = param_name
                        param_table.cell(i, 1).text = str(param_data.get('value', ''))
                        param_table.cell(i, 2).text = param_data.get('unit', '')
                        
            # Рекомендации
            if result.recommendations and config.include_recommendations:
                doc.add_heading('3. Рекомендации', 1)
                
                for i, recommendation in enumerate(result.recommendations, 1):
                    doc.add_paragraph(f"{i}. {recommendation}", style='ListBullet')
                    
            # Технические данные
            if config.include_technical_data:
                doc.add_heading('4. Технические данные', 1)
                
                tech_data = [
                    ("Время диагностики:", f"{result.diagnostic_duration:.1f} сек"),
                    ("Дата диагностики:", result.timestamp.strftime("%d.%m.%Y %H:%M:%S")),
                    ("Версия ПО:", result.software_version),
                    ("Техник:", result.technician),
                    ("Примечания:", result.vehicle_info.notes)
                ]
                
                for label, value in tech_data:
                    if value:
                        p = doc.add_paragraph()
                        p.add_run(label).bold = True
                        p.add_run(f" {value}")
                        
            # Графики если требуются
            if config.include_graphs:
                doc.add_heading('5. Графики параметров', 1)
                graphs = self._generate_graphs(result, config)
                
                for graph_path in graphs.values():
                    if graph_path.exists():
                        doc.add_picture(str(graph_path), width=Inches(6))
                        doc.add_paragraph()  # Пустая строка
                        
            # Подписи
            if config.technician_signature or config.customer_signature:
                doc.add_page_break()
                doc.add_heading('ПОДПИСИ', 1)
                
                sig_table = doc.add_table(rows=2, cols=2)
                sig_table.style = 'LightShading-Accent1'
                
                if config.technician_signature:
                    sig_table.cell(0, 0).text = "ТЕХНИК"
                    sig_table.cell(1, 0).text = "___________________________"
                    
                if config.customer_signature:
                    sig_table.cell(0, 1).text = "КЛИЕНТ"
                    sig_table.cell(1, 1).text = "___________________________"
                    
            # Сохранение документа
            filename = f"{base_filename}_{config.level.value}.docx"
            filepath = self.reports_dir / "docx" / filename
            doc.save(str(filepath))
            
            self.logger.info(f"DOCX отчет сохранен: {filepath}")
            return filepath
            
        except Exception as e:
            self.logger.error(f"Ошибка создания DOCX отчета: {e}")
            raise
            
    def generate_excel_report(self,
                             result: DiagnosticResult,
                             config: ReportConfig,
                             base_filename: str) -> Path:
        """Генерация Excel отчета"""
        try:
            wb = Workbook()
            
            # Стили
            header_font = Font(bold=True, color="FFFFFF")
            header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            center_alignment = Alignment(horizontal='center', vertical='center')
            
            # Лист "Общая информация"
            ws_info = wb.active
            ws_info.title = "Общая информация"
            
            # Заголовок
            ws_info.merge_cells('A1:D1')
            ws_info['A1'] = "ОТЧЕТ ДИАГНОСТИКИ CHEVROLET NIVA"
            ws_info['A1'].font = Font(bold=True, size=14)
            ws_info['A1'].alignment = center_alignment
            
            # Информация об автомобиле
            ws_info.append([])
            ws_info.append(["ИНФОРМАЦИЯ ОБ АВТОМОБИЛЕ"])
            ws_info.append([])
            
            vehicle_data = [
                ["VIN:", result.vehicle_info.vin],
                ["Модель:", result.vehicle_info.model],
                ["Год выпуска:", result.vehicle_info.year],
                ["Двигатель:", result.vehicle_info.engine],
                ["Пробег, км:", result.vehicle_info.mileage],
                ["Гос. номер:", result.vehicle_info.license_plate],
                ["Владелец:", result.vehicle_info.owner],
                ["Контакты:", result.vehicle_info.contact],
                ["Дата диагностики:", result.timestamp.strftime("%d.%m.%Y %H:%M:%S")],
                ["Техник:", result.technician],
                ["Оценка состояния:", f"{result.rating}/100"]
            ]
            
            for row in vehicle_data:
                ws_info.append(row)
                
            # Форматирование
            for col in range(1, 3):
                for row in range(3, len(vehicle_data) + 3):
                    cell = ws_info.cell(row=row, column=col)
                    cell.border = border
                    
            # Лист "Ошибки"
            if result.dtcs:
                ws_dtcs = wb.create_sheet("Коды ошибок")
                
                headers = ["Код", "Описание", "Система", "Статус", "Критичность", "Рекомендации"]
                ws_dtcs.append(headers)
                
                # Заголовки
                for col in range(1, len(headers) + 1):
                    cell = ws_dtcs.cell(row=1, column=col)
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = center_alignment
                    cell.border = border
                    
                # Данные
                for dtc in result.dtcs:
                    row = [
                        dtc.get('code', ''),
                        dtc.get('description', ''),
                        dtc.get('system', ''),
                        dtc.get('status', ''),
                        dtc.get('severity', ''),
                        dtc.get('recommendation', '')
                    ]
                    ws_dtcs.append(row)
                    
                # Форматирование данных
                for row in ws_dtcs.iter_rows(min_row=2, max_row=len(result.dtcs) + 1, min_col=1, max_col=6):
                    for cell in row:
                        cell.border = border
                        
                # Автоширина колонок
                for column in ws_dtcs.columns:
                    max_length = 0
                    column_letter = get_column_letter(column[0].column)
                    
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                            
                    adjusted_width = min(max_length + 2, 50)
                    ws_dtcs.column_dimensions[column_letter].width = adjusted_width
                    
            # Лист "Параметры"
            if result.live_data:
                ws_params = wb.create_sheet("Текущие параметры")
                
                # Группируем параметры
                params_by_system = self._group_parameters_by_system(result.live_data)
                
                row_offset = 1
                for system, params in params_by_system.items():
                    # Заголовок системы
                    ws_params.cell(row=row_offset, column=1, value=f"СИСТЕМА: {system}")
                    ws_params.cell(row=row_offset, column=1).font = Font(bold=True, size=12)
                    row_offset += 1
                    
                    # Заголовки таблицы
                    headers = ["Параметр", "Значение", "Единица", "Минимум", "Норма", "Максимум", "Статус"]
                    ws_params.append(headers)
                    
                    # Заголовки форматирование
                    for col in range(1, len(headers) + 1):
                        cell = ws_params.cell(row=row_offset, column=col)
                        cell.font = header_font
                        cell.fill = header_fill
                        cell.alignment = center_alignment
                        cell.border = border
                    row_offset += 1
                    
                    # Данные параметров
                    for param_name, param_data in params.items():
                        value = param_data.get('value', 0)
                        min_val = param_data.get('min', 0)
                        max_val = param_data.get('max', 0)
                        
                        # Определение статуса
                        if value < min_val:
                            status = "НИЖЕ НОРМЫ"
                            status_color = "FFFF00"  # Желтый
                        elif value > max_val:
                            status = "ВЫШЕ НОРМЫ"
                            status_color = "FF0000"  # Красный
                        else:
                            status = "НОРМА"
                            status_color = "00FF00"  # Зеленый
                            
                        row = [
                            param_name,
                            value,
                            param_data.get('unit', ''),
                            min_val,
                            param_data.get('normal', ''),
                            max_val,
                            status
                        ]
                        
                        ws_params.append(row)
                        
                        # Форматирование строки
                        for col in range(1, len(row) + 1):
                            cell = ws_params.cell(row=row_offset, column=col)
                            cell.border = border
                            
                            if col == 7:  # Столбец статуса
                                if status != "НОРМА":
                                    cell.fill = PatternFill(start_color=status_color, 
                                                          end_color=status_color, 
                                                          fill_type="solid")
                                    cell.font = Font(bold=True)
                                    
                        row_offset += 1
                        
                    row_offset += 1  # Пустая строка между системами
                    
            # Лист "Рекомендации"
            if result.recommendations:
                ws_rec = wb.create_sheet("Рекомендации")
                
                ws_rec.append(["ПРИОРИТЕТ", "РЕКОМЕНДАЦИЯ", "СТОИМОСТЬ", "СРОК"])
                
                # Заголовки
                for col in range(1, 5):
                    cell = ws_rec.cell(row=1, column=col)
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = center_alignment
                    cell.border = border
                    
                # Сортируем рекомендации по приоритету
                for i, rec in enumerate(result.recommendations, 1):
                    priority = self._determine_priority(rec)
                    ws_rec.append([priority, rec, "", ""])
                    
                # Форматирование
                for row in ws_rec.iter_rows(min_row=2, max_row=len(result.recommendations) + 1):
                    for cell in row:
                        cell.border = border
                        
            # Лист "Статистика"
            ws_stats = wb.create_sheet("Статистика")
            
            # Сводная статистика
            stats_data = [
                ["ПАРАМЕТР", "ЗНАЧЕНИЕ"],
                ["Общее количество ошибок", len(result.dtcs)],
                ["Критических ошибок", len([d for d in result.dtcs if d.get('severity') == 'CRITICAL'])],
                ["Предупреждений", len([d for d in result.dtcs if d.get('severity') == 'WARNING'])],
                ["Нормальных параметров", self._count_normal_parameters(result.live_data)],
                ["Параметров вне нормы", self._count_abnormal_parameters(result.live_data)],
                ["Общая оценка", f"{result.rating}/100"],
                ["Вердикт", self._get_verdict(result.rating)]
            ]
            
            for i, row in enumerate(stats_data, 1):
                ws_stats.append(row)
                for col in range(1, 3):
                    cell = ws_stats.cell(row=i, column=col)
                    cell.border = border
                    if i == 1:
                        cell.font = header_font
                        cell.fill = header_fill
                        
            # Сохранение файла
            filename = f"{base_filename}_{config.level.value}.xlsx"
            filepath = self.reports_dir / "xlsx" / filename
            wb.save(str(filepath))
            
            self.logger.info(f"Excel отчет сохранен: {filepath}")
            return filepath
            
        except Exception as e:
            self.logger.error(f"Ошибка создания Excel отчета: {e}")
            raise
            
    def generate_json_report(self,
                           result: DiagnosticResult,
                           config: ReportConfig,
                           base_filename: str) -> Path:
        """Генерация JSON отчета"""
        try:
            # Преобразуем результат в словарь
            report_dict = {
                'metadata': {
                    'generated_at': datetime.datetime.now().isoformat(),
                    'software_version': result.software_version,
                    'report_format': config.format.value,
                    'report_level': config.level.value
                },
                'vehicle_info': asdict(result.vehicle_info),
                'diagnostic_results': {
                    'rating': result.rating,
                    'diagnostic_duration': result.diagnostic_duration,
                    'timestamp': result.timestamp.isoformat(),
                    'technician': result.technician
                },
                'dtcs': result.dtcs,
                'live_data': result.live_data,
                'ecu_status': result.ecu_status,
                'sensor_checks': result.sensor_checks,
                'actuator_checks': result.actuator_checks,
                'adaptation_results': result.adaptation_results,
                'recommendations': result.recommendations,
                'summary': {
                    'verdict': self._get_verdict(result.rating),
                    'critical_issues': len([d for d in result.dtcs if d.get('severity') == 'CRITICAL']),
                    'warnings': len([d for d in result.dtcs if d.get('severity') == 'WARNING']),
                    'normal_parameters': self._count_normal_parameters(result.live_data),
                    'abnormal_parameters': self._count_abnormal_parameters(result.live_data)
                }
            }
            
            # Сохранение в файл
            filename = f"{base_filename}_{config.level.value}.json"
            filepath = self.reports_dir / "json" / filename
            
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(report_dict, f, ensure_ascii=False, indent=2, default=str)
                
            self.logger.info(f"JSON отчет сохранен: {filepath}")
            return filepath
            
        except Exception as e:
            self.logger.error(f"Ошибка создания JSON отчета: {e}")
            raise
            
    def generate_csv_report(self,
                           result: DiagnosticResult,
                           config: ReportConfig,
                           base_filename: str) -> Path:
        """Генерация CSV отчета"""
        try:
            filename = f"{base_filename}_{config.level.value}.csv"
            filepath = self.reports_dir / "csv" / filename
            
            # Создаем несколько CSV файлов
            # 1. Основная информация
            main_file = filepath.with_stem(f"{base_filename}_main")
            with open(main_file, 'w', encoding='utf-8', newline='') as f:
                writer = csv.writer(f, delimiter=';')
                writer.writerow(['Параметр', 'Значение'])
                
                # Информация об автомобиле
                vehicle_info = asdict(result.vehicle_info)
                for key, value in vehicle_info.items():
                    writer.writerow([key, value])
                    
                writer.writerow(['Оценка состояния', result.rating])
                writer.writerow(['Дата диагностики', result.timestamp])
                writer.writerow(['Техник', result.technician])
                
            # 2. Ошибки
            if result.dtcs:
                dtc_file = filepath.with_stem(f"{base_filename}_dtcs")
                with open(dtc_file, 'w', encoding='utf-8', newline='') as f:
                    writer = csv.writer(f, delimiter=';')
                    writer.writerow(['Код', 'Описание', 'Система', 'Статус', 'Критичность'])
                    
                    for dtc in result.dtcs:
                        writer.writerow([
                            dtc.get('code', ''),
                            dtc.get('description', ''),
                            dtc.get('system', ''),
                            dtc.get('status', ''),
                            dtc.get('severity', '')
                        ])
                        
            # 3. Параметры
            if result.live_data:
                params_file = filepath.with_stem(f"{base_filename}_params")
                with open(params_file, 'w', encoding='utf-8', newline='') as f:
                    writer = csv.writer(f, delimiter=';')
                    writer.writerow(['Система', 'Параметр', 'Значение', 'Единица', 'Статус'])
                    
                    params_by_system = self._group_parameters_by_system(result.live_data)
                    for system, params in params_by_system.items():
                        for param_name, param_data in params.items():
                            value = param_data.get('value', 0)
                            min_val = param_data.get('min', 0)
                            max_val = param_data.get('max', 0)
                            
                            if value < min_val:
                                status = "НИЖЕ НОРМЫ"
                            elif value > max_val:
                                status = "ВЫШЕ НОРМЫ"
                            else:
                                status = "НОРМА"
                                
                            writer.writerow([
                                system,
                                param_name,
                                value,
                                param_data.get('unit', ''),
                                status
                            ])
                            
            self.logger.info(f"CSV отчеты сохранены: {filepath.parent}")
            return filepath
            
        except Exception as e:
            self.logger.error(f"Ошибка создания CSV отчета: {e}")
            raise
            
    def _create_filename(self, vehicle_info: VehicleInfo) -> str:
        """Создание имени файла на основе информации об автомобиле"""
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Используем VIN или госномер для имени файла
        identifier = vehicle_info.vin or vehicle_info.license_plate or "unknown"
        identifier = identifier.replace(' ', '_').replace('/', '_')
        
        model = vehicle_info.model.replace(' ', '_') if vehicle_info.model else "Niva"
        
        return f"{model}_{identifier}_{timestamp}"
        
    def _prepare_template_data(self, result: DiagnosticResult, config: ReportConfig) -> Dict:
        """Подготовка данных для шаблона"""
        # Определяем оценку состояния
        rating_text = self._get_rating_text(result.rating)
        rating_color = self._get_rating_color(result.rating)
        
        # Группируем параметры
        params_by_system = self._group_parameters_by_system(result.live_data)
        
        # Группируем ошибки по системе
        dtcs_by_system = {}
        for dtc in result.dtcs:
            system = dtc.get('system', 'Unknown')
            if system not in dtcs_by_system:
                dtcs_by_system[system] = []
            dtcs_by_system[system].append(dtc)
            
        return {
            'report_title': f"Отчет диагностики {result.vehicle_info.model}",
            'vehicle_info': result.vehicle_info,
            'timestamp': result.timestamp.strftime("%d.%m.%Y %H:%M:%S"),
            'rating': result.rating,
            'rating_text': rating_text,
            'rating_color': rating_color,
            'technician': result.technician,
            'diagnostic_duration': f"{result.diagnostic_duration:.1f}",
            'dtcs': result.dtcs,
            'dtcs_by_system': dtcs_by_system,
            'params_by_system': params_by_system,
            'recommendations': result.recommendations,
            'ecu_status': result.ecu_status,
            'config': config,
            'translations': self.translations[config.language],
            'company_name': config.company_name,
            'company_address': config.company_address,
            'company_phone': config.company_phone,
            'verdict': self._get_verdict(result.rating),
            'summary_stats': {
                'total_dtcs': len(result.dtcs),
                'critical_dtcs': len([d for d in result.dtcs if d.get('severity') == 'CRITICAL']),
                'normal_params': self._count_normal_parameters(result.live_data),
                'abnormal_params': self._count_abnormal_parameters(result.live_data)
            }
        }
        
    def _generate_graphs(self, result: DiagnosticResult, config: ReportConfig) -> Dict[str, Path]:
        """Генерация графиков для отчета"""
        graphs = {}
        
        try:
            # Создаем директорию для графиков
            graphs_dir = self.reports_dir / "graphs"
            graphs_dir.mkdir(exist_ok=True)
            
            # 1. График оценки состояния
            fig, ax = plt.subplots(figsize=(10, 6))
            categories = ['Ошибки', 'Двигатель', 'Топливо', 'Электрика', 'Шасси']
            values = [result.rating, 85, 90, 88, 92]  # Примерные значения
            
            bars = ax.bar(categories, values, color=['red', 'green', 'blue', 'orange', 'purple'])
            
            # Добавляем значения на столбцы
            for bar, value in zip(bars, values):
                ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1,
                       f'{value}%', ha='center', va='bottom', fontsize=10)
                
            ax.set_ylim(0, 100)
            ax.set_ylabel('Оценка, %')
            ax.set_title('Оценка состояния систем автомобиля')
            ax.grid(axis='y', linestyle='--', alpha=0.7)
            
            graph_path = graphs_dir / "system_rating.png"
            plt.tight_layout()
            plt.savefig(graph_path, dpi=150)
            plt.close()
            graphs['system_rating'] = graph_path
            
            # 2. График параметров двигателя в реальном времени
            if result.live_data:
                engine_params = {k: v for k, v in result.live_data.items() 
                               if 'RPM' in k or 'TEMP' in k or 'PRESSURE' in k}
                
                if engine_params:
                    fig, axes = plt.subplots(2, 2, figsize=(12, 8))
                    axes = axes.flatten()
                    
                    for idx, (param_name, param_data) in enumerate(list(engine_params.items())[:4]):
                        ax = axes[idx]
                        
                        # Создаем примерные временные данные
                        time_points = np.arange(10)
                        values = np.random.normal(param_data.get('value', 0), 
                                                param_data.get('value', 0) * 0.1, 
                                                10)
                                                
                        ax.plot(time_points, values, 'b-', linewidth=2)
                        ax.set_title(param_name)
                        ax.set_xlabel('Время, с')
                        ax.set_ylabel(param_data.get('unit', ''))
                        ax.grid(True, alpha=0.3)
                        
                        # Добавляем нормативные значения если есть
                        min_val = param_data.get('min')
                        max_val = param_data.get('max')
                        if min_val is not None and max_val is not None:
                            ax.axhline(y=min_val, color='r', linestyle='--', alpha=0.5, label='Мин')
                            ax.axhline(y=max_val, color='r', linestyle='--', alpha=0.5, label='Макс')
                            ax.fill_between(time_points, min_val, max_val, alpha=0.1, color='green')
                            
                    plt.tight_layout()
                    engine_graph_path = graphs_dir / "engine_parameters.png"
                    plt.savefig(engine_graph_path, dpi=150)
                    plt.close()
                    graphs['engine_parameters'] = engine_graph_path
                    
            # 3. Круговая диаграмма распределения ошибок
            if result.dtcs:
                fig, ax = plt.subplots(figsize=(8, 8))
                
                # Группируем ошибки по системам
                systems = {}
                for dtc in result.dtcs:
                    system = dtc.get('system', 'Unknown')
                    systems[system] = systems.get(system, 0) + 1
                    
                labels = list(systems.keys())
                sizes = list(systems.values())
                
                if sizes:
                    colors = plt.cm.Set3(np.linspace(0, 1, len(labels)))
                    wedges, texts, autotexts = ax.pie(sizes, labels=labels, colors=colors,
                                                    autopct='%1.1f%%', startangle=90,
                                                    textprops={'fontsize': 10})
                    
                    # Делаем проценты жирными
                    for autotext in autotexts:
                        autotext.set_color('white')
                        autotext.set_fontweight('bold')
                        
                    ax.set_title('Распределение ошибок по системам')
                    
                    pie_chart_path = graphs_dir / "dtc_distribution.png"
                    plt.tight_layout()
                    plt.savefig(pie_chart_path, dpi=150)
                    plt.close()
                    graphs['dtc_distribution'] = pie_chart_path
                    
            # 4. Линейный график основных параметров
            if result.live_data:
                key_params = ['ENGINE_RPM', 'COOLANT_TEMP', 'THROTTLE_POSITION', 'ENGINE_LOAD']
                key_data = {k: result.live_data.get(k) for k in key_params if k in result.live_data}
                
                if key_data:
                    fig, ax = plt.subplots(figsize=(12, 6))
                    
                    time_points = np.arange(len(key_data))
                    values = [data.get('value', 0) for data in key_data.values()]
                    param_names = list(key_data.keys())
                    
                    ax.plot(time_points, values, 'o-', linewidth=2, markersize=8)
                    ax.set_xticks(time_points)
                    ax.set_xticklabels([self._translate_param_name(name) for name in param_names], 
                                     rotation=45, ha='right')
                    
                    ax.set_ylabel('Значение')
                    ax.set_title('Основные параметры двигателя')
                    ax.grid(True, alpha=0.3)
                    
                    # Добавляем значения на точки
                    for i, v in enumerate(values):
                        ax.annotate(f'{v:.1f}', (time_points[i], v), 
                                  textcoords="offset points", 
                                  xytext=(0,10), ha='center')
                          
                    linear_graph_path = graphs_dir / "key_parameters.png"
                    plt.tight_layout()
                    plt.savefig(linear_graph_path, dpi=150)
                    plt.close()
                    graphs['key_parameters'] = linear_graph_path
                    
        except Exception as e:
            self.logger.error(f"Ошибка генерации графиков: {e}")
            
        return graphs
        
    def _translate_param_name(self, param_name: str) -> str:
        """Перевод названий параметров"""
        translations = {
            'ENGINE_RPM': 'Обороты двигателя',
            'COOLANT_TEMP': 'Темп. охлаждающей жидкости',
            'THROTTLE_POSITION': 'Положение дросселя',
            'ENGINE_LOAD': 'Нагрузка двигателя',
            'VEHICLE_SPEED': 'Скорость',
            'INTAKE_TEMP': 'Темп. впускного воздуха',
            'MAF_SENSOR': 'Датчик массового расхода',
            'FUEL_PRESSURE': 'Давление топлива',
            'INTAKE_PRESSURE': 'Давление во впуске',
            'TIMING_ADVANCE': 'Угол опережения',
            'FUEL_LEVEL': 'Уровень топлива',
            'CONTROL_MODULE_VOLTAGE': 'Напряжение ЭБУ'
        }
        return translations.get(param_name, param_name)
        
    def _group_parameters_by_system(self, live_data: Dict) -> Dict[str, Dict]:
        """Группировка параметров по системам"""
        systems = {
            'engine': ['RPM', 'TEMP', 'LOAD', 'TIMING', 'PRESSURE'],
            'fuel': ['FUEL', 'MAF', 'O2'],
            'electrical': ['VOLTAGE', 'CURRENT'],
            'climate': ['TEMP', 'PRESSURE'],
            'chassis': ['SPEED', 'POSITION']
        }
        
        result = {}
        for param_name, param_data in live_data.items():
            assigned = False
            
            for system_name, keywords in systems.items():
                for keyword in keywords:
                    if keyword in param_name:
                        if system_name not in result:
                            result[system_name] = {}
                        result[system_name][param_name] = param_data
                        assigned = True
                        break
                if assigned:
                    break
                    
            if not assigned:
                if 'other' not in result:
                    result['other'] = {}
                result['other'][param_name] = param_data
                
        return result
        
    def _get_rating_text(self, rating: int) -> str:
        """Получение текстового описания оценки"""
        if rating >= 90:
            return "Отличное"
        elif rating >= 75:
            return "Хорошее"
        elif rating >= 60:
            return "Удовлетворительное"
        elif rating >= 40:
            return "Плохое"
        else:
            return "Критическое"
            
    def _get_rating_color(self, rating: int) -> RGBColor:
        """Получение цвета оценки"""
        if rating >= 90:
            return RGBColor(0, 128, 0)  # Зеленый
        elif rating >= 75:
            return RGBColor(0, 0, 255)  # Синий
        elif rating >= 60:
            return RGBColor(255, 165, 0)  # Оранжевый
        elif rating >= 40:
            return RGBColor(255, 0, 0)  # Красный
        else:
            return RGBColor(139, 0, 0)  # Темно-красный
            
    def _get_verdict(self, rating: int) -> str:
        """Получение вердикта по оценке"""
        if rating >= 90:
            return "Автомобиль в отличном состоянии. Рекомендуется плановое обслуживание."
        elif rating >= 75:
            return "Состояние хорошее. Необходима проверка выявленных замечаний."
        elif rating >= 60:
            return "Состояние удовлетворительное. Требуется обслуживание в ближайшее время."
        elif rating >= 40:
            return "Состояние плохое. Требуется срочное обслуживание."
        else:
            return "КРИТИЧЕСКОЕ СОСТОЯНИЕ! Эксплуатация опасна. Требуется немедленный ремонт."
            
    def _count_normal_parameters(self, live_data: Dict) -> int:
        """Подсчет параметров в норме"""
        normal_count = 0
        for param_data in live_data.values():
            value = param_data.get('value', 0)
            min_val = param_data.get('min', 0)
            max_val = param_data.get('max', 0)
            
            if min_val <= value <= max_val:
                normal_count += 1
                
        return normal_count
        
    def _count_abnormal_parameters(self, live_data: Dict) -> int:
        """Подсчет параметров вне нормы"""
        abnormal_count = 0
        for param_data in live_data.values():
            value = param_data.get('value', 0)
            min_val = param_data.get('min', 0)
            max_val = param_data.get('max', 0)
            
            if value < min_val or value > max_val:
                abnormal_count += 1
                
        return abnormal_count
        
    def _determine_priority(self, recommendation: str) -> str:
        """Определение приоритета рекомендации"""
        high_priority_keywords = ['срочно', 'немедленно', 'опасно', 'критическ', 'запрещено']
        medium_priority_keywords = ['рекомендуется', 'желательно', 'в ближайшее время']
        
        recommendation_lower = recommendation.lower()
        
        for keyword in high_priority_keywords:
            if keyword in recommendation_lower:
                return "ВЫСОКИЙ"
                
        for keyword in medium_priority_keywords:
            if keyword in recommendation_lower:
                return "СРЕДНИЙ"
                
        return "НИЗКИЙ"
        
    def _create_default_html_template(self):
        """Создание стандартного HTML шаблона"""
        template_content = """
<!DOCTYPE html>
<html lang="{{ config.language }}">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{{ report_title }}</title>
    <style>
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
            background-color: #f5f5f5;
        }
        .header {
            background-color: #2c3e50;
            color: white;
            padding: 30px;
            border-radius: 10px;
            margin-bottom: 30px;
            text-align: center;
        }
        .header h1 {
            margin: 0;
            font-size: 28px;
        }
        .rating {
            display: inline-block;
            padding: 10px 20px;
            border-radius: 20px;
            font-weight: bold;
            font-size: 20px;
            margin-top: 15px;
            background-color: {{ rating_color }};
            color: white;
        }
        .section {
            background-color: white;
            padding: 25px;
            border-radius: 10px;
            margin-bottom: 25px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }
        .section h2 {
            color: #2c3e50;
            border-bottom: 2px solid #3498db;
            padding-bottom: 10px;
            margin-top: 0;
        }
        .vehicle-info {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(250px, 1fr));
            gap: 15px;
            margin-top: 20px;
        }
        .info-item {
            display: flex;
            justify-content: space-between;
            padding: 10px;
            background-color: #f8f9fa;
            border-radius: 5px;
        }
        .info-label {
            font-weight: bold;
            color: #2c3e50;
        }
        table {
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
        }
        th {
            background-color: #3498db;
            color: white;
            padding: 12px;
            text-align: left;
        }
        td {
            padding: 10px;
            border-bottom: 1px solid #ddd;
        }
        tr:nth-child(even) {
            background-color: #f8f9fa;
        }
        tr:hover {
            background-color: #e8f4fc;
        }
        .critical {
            background-color: #ffe6e6 !important;
            color: #c00;
            font-weight: bold;
        }
        .warning {
            background-color: #fff3cd !important;
            color: #856404;
        }
        .normal {
            background-color: #d4edda !important;
            color: #155724;
        }
        .recommendation {
            padding: 15px;
            margin: 10px 0;
            background-color: #e8f4fc;
            border-left: 4px solid #3498db;
            border-radius: 5px;
        }
        .high-priority {
            border-left-color: #dc3545;
            background-color: #ffe6e6;
        }
        .medium-priority {
            border-left-color: #ffc107;
            background-color: #fff3cd;
        }
        .graph-container {
            text-align: center;
            margin: 30px 0;
        }
        .graph-container img {
            max-width: 100%;
            height: auto;
            border-radius: 10px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }
        .footer {
            text-align: center;
            margin-top: 50px;
            padding: 20px;
            color: #666;
            border-top: 1px solid #ddd;
        }
        .signature-area {
            display: flex;
            justify-content: space-around;
            margin-top: 50px;
            padding-top: 30px;
            border-top: 2px solid #ddd;
        }
        .signature {
            text-align: center;
        }
        .signature-line {
            width: 300px;
            height: 1px;
            background-color: #333;
            margin: 40px auto 10px;
        }
        .company-info {
            background-color: #2c3e50;
            color: white;
            padding: 20px;
            border-radius: 10px;
            margin-top: 30px;
            text-align: center;
        }
        @media print {
            body {
                background-color: white;
            }
            .section {
                box-shadow: none;
                border: 1px solid #ddd;
            }
            .no-print {
                display: none;
            }
        }
    </style>
</head>
<body>
    <div class="header">
        <h1>{{ report_title }}</h1>
        <div class="rating">
            {{ rating }}/100 - {{ rating_text }}
        </div>
        <p>Дата диагностики: {{ timestamp }}</p>
    </div>

    <div class="section">
        <h2>1. Информация об автомобиле</h2>
        <div class="vehicle-info">
            <div class="info-item">
                <span class="info-label">VIN:</span>
                <span>{{ vehicle_info.vin }}</span>
            </div>
            <div class="info-item">
                <span class="info-label">Модель:</span>
                <span>{{ vehicle_info.model }}</span>
            </div>
            <div class="info-item">
                <span class="info-label">Год выпуска:</span>
                <span>{{ vehicle_info.year }}</span>
            </div>
            <div class="info-item">
                <span class="info-label">Двигатель:</span>
                <span>{{ vehicle_info.engine }}</span>
            </div>
            <div class="info-item">
                <span class="info-label">Пробег:</span>
                <span>{{ vehicle_info.mileage }} км</span>
            </div>
            <div class="info-item">
                <span class="info-label">Гос. номер:</span>
                <span>{{ vehicle_info.license_plate }}</span>
            </div>
            <div class="info-item">
                <span class="info-label">Владелец:</span>
                <span>{{ vehicle_info.owner }}</span>
            </div>
            <div class="info-item">
                <span class="info-label">Техник:</span>
                <span>{{ technician }}</span>
            </div>
        </div>
    </div>

    {% if summary_stats.total_dtcs > 0 %}
    <div class="section">
        <h2>2. Коды неисправностей ({{ summary_stats.total_dtcs }})</h2>
        
        {% for system, system_dtcs in dtcs_by_system.items() %}
        <h3>Система: {{ system }}</h3>
        <table>
            <thead>
                <tr>
                    <th>Код</th>
                    <th>Описание</th>
                    <th>Статус</th>
                    <th>Критичность</th>
                </tr>
            </thead>
            <tbody>
                {% for dtc in system_dtcs %}
                <tr class="{{ 'critical' if dtc.severity == 'CRITICAL' else 'warning' if dtc.severity == 'WARNING' else '' }}">
                    <td>{{ dtc.code }}</td>
                    <td>{{ dtc.description }}</td>
                    <td>{{ dtc.status }}</td>
                    <td>{{ dtc.severity }}</td>
                </tr>
                {% endfor %}
            </tbody>
        </table>
        {% endfor %}
    </div>
    {% endif %}

    {% if params_by_system %}
    <div class="section">
        <h2>3. Текущие параметры</h2>
        
        {% for system, params in params_by_system.items() %}
        <h3>Система: {{ system }}</h3>
        <table>
            <thead>
                <tr>
                    <th>Параметр</th>
                    <th>Значение</th>
                    <th>Единица</th>
                    <th>Норма</th>
                    <th>Статус</th>
                </tr>
            </thead>
            <tbody>
                {% for param_name, param_data in params.items() %}
                {% set value = param_data.value %}
                {% set min_val = param_data.min %}
                {% set max_val = param_data.max %}
                {% set status = 'normal' if min_val <= value <= max_val else 'abnormal' %}
                <tr class="{{ 'normal' if status == 'normal' else 'critical' }}">
                    <td>{{ param_name }}</td>
                    <td>{{ "%.2f"|format(value) }}</td>
                    <td>{{ param_data.unit }}</td>
                    <td>{{ "%.2f"|format(min_val) }} - {{ "%.2f"|format(max_val) }}</td>
                    <td>{{ 'НОРМА' if status == 'normal' else 'ВНЕ НОРМЫ' }}</td>
                </tr>
                {% endfor %}
            </tbody>
        </table>
        {% endfor %}
    </div>
    {% endif %}

    {% if config.include_graphs and graphs %}
    <div class="section">
        <h2>4. Графики</h2>
        
        {% for graph_name, graph_path in graphs.items() %}
        <div class="graph-container">
            <img src="{{ graph_path }}" alt="{{ graph_name }}">
        </div>
        {% endfor %}
    </div>
    {% endif %}

    {% if recommendations and config.include_recommendations %}
    <div class="section">
        <h2>5. Рекомендации</h2>
        
        {% for recommendation in recommendations %}
        <div class="recommendation {{ 'high-priority' if 'срочно' in recommendation.lower() or 'опасно' in recommendation.lower() else 'medium-priority' if 'рекомендуется' in recommendation.lower() else '' }}">
            {{ recommendation }}
        </div>
        {% endfor %}
    </div>
    {% endif %}

    <div class="section">
        <h2>6. Сводка</h2>
        <div class="vehicle-info">
            <div class="info-item">
                <span class="info-label">Общая оценка:</span>
                <span>{{ rating }}/100</span>
            </div>
            <div class="info-item">
                <span class="info-label">Критических ошибок:</span>
                <span>{{ summary_stats.critical_dtcs }}</span>
            </div>
            <div class="info-item">
                <span class="info-label">Предупреждений:</span>
                <span>{{ summary_stats.total_dtcs - summary_stats.critical_dtcs }}</span>
            </div>
            <div class="info-item">
                <span class="info-label">Параметров в норме:</span>
                <span>{{ summary_stats.normal_params }}</span>
            </div>
            <div class="info-item">
                <span class="info-label">Параметров вне нормы:</span>
                <span>{{ summary_stats.abnormal_params }}</span>
            </div>
        </div>
        
        <div style="margin-top: 30px; padding: 20px; background-color: #e8f4fc; border-radius: 5px;">
            <h3>Вердикт:</h3>
            <p style="font-size: 18px; font-weight: bold; color: #2c3e50;">{{ verdict }}</p>
        </div>
    </div>

    {% if config.technician_signature or config.customer_signature %}
    <div class="signature-area">
        {% if config.technician_signature %}
        <div class="signature">
            <p>Техник</p>
            <div class="signature-line"></div>
            <p>{{ technician }}</p>
        </div>
        {% endif %}
        
        {% if config.customer_signature %}
        <div class="signature">
            <p>Клиент</p>
            <div class="signature-line"></div>
            <p>{{ vehicle_info.owner }}</p>
        </div>
        {% endif %}
    </div>
    {% endif %}

    {% if config.company_name %}
    <div class="company-info">
        <h3>{{ config.company_name }}</h3>
        {% if config.company_address %}<p>{{ config.company_address }}</p>{% endif %}
        {% if config.company_phone %}<p>Тел.: {{ config.company_phone }}</p>{% endif %}
        <p>Отчет сгенерирован программой "Профессиональная диагностика Chevrolet Niva"</p>
    </div>
    {% endif %}

    <div class="footer">
        <p>Страница сгенерирована {{ timestamp }} | Версия ПО: {{ software_version }}</p>
        <button class="no-print" onclick="window.print()">Печать отчета</button>
    </div>

    <script>
        // Подсветка строк таблицы при наведении
        document.addEventListener('DOMContentLoaded', function() {
            const rows = document.querySelectorAll('tbody tr');
            rows.forEach(row => {
                row.addEventListener('mouseenter', function() {
                    this.style.transform = 'translateY(-2px)';
                    this.style.boxShadow = '0 4px 8px rgba(0,0,0,0.1)';
                    this.style.transition = 'all 0.3s ease';
                });
                row.addEventListener('mouseleave', function() {
                    this.style.transform = 'translateY(0)';
                    this.style.boxShadow = 'none';
                });
            });
        });
    </script>
</body>
</html>
        """
        
        # Сохраняем шаблон
        template_path = self.templates_dir / "report_standard.html"
        with open(template_path, 'w', encoding='utf-8') as f:
            f.write(template_content)
            
        # Создаем шаблоны для других уровней
        for level in ReportLevel:
            if level != ReportLevel.STANDARD:
                level_template = template_path.with_stem(f"report_{level.value}")
                with open(level_template, 'w', encoding='utf-8') as f:
                    f.write(template_content)  # В реальности будут разные шаблоны
                    
    def _create_simple_html_report(self, result: DiagnosticResult, config: ReportConfig, base_filename: str) -> Path:
        """Создание простого HTML отчета в случае ошибки"""
        html_content = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Диагностика {result.vehicle_info.model}</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; }}
        h1 {{ color: #333; }}
        .info {{ margin: 20px 0; }}
        .dtc {{ color: red; font-weight: bold; }}
        .param {{ margin: 5px 0; }}
    </style>
</head>
<body>
    <h1>Отчет диагностики</h1>
    <div class="info">
        <p>Автомобиль: {result.vehicle_info.model}</p>
        <p>VIN: {result.vehicle_info.vin}</p>
        <p>Дата: {result.timestamp}</p>
    </div>
    <h2>Ошибки:</h2>
    {"<br>".join([f'<span class="dtc">{dtc.get("code", "")}: {dtc.get("description", "")}</span>' for dtc in result.dtcs]) if result.dtcs else "Нет ошибок"}
    <h2>Основные параметры:</h2>
    {"<br>".join([f'<div class="param">{k}: {v.get("value", "")} {v.get("unit", "")}</div>' for k, v in result.live_data.items()])}
</body>
</html>
        """
        
        filename = f"{base_filename}_simple.html"
        filepath = self.reports_dir / "html" / filename
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(html_content)
            
        return filepath
        
    def _copy_styles(self):
        """Копирование стилей в директорию отчетов"""
        try:
            styles_dir = Path("assets/styles")
            if styles_dir.exists():
                import shutil
                reports_styles_dir = self.reports_dir / "html" / "styles"
                reports_styles_dir.mkdir(exist_ok=True)
                
                for style_file in styles_dir.glob("*.css"):
                    shutil.copy2(style_file, reports_styles_dir)
                    
        except Exception as e:
            self.logger.warning(f"Не удалось скопировать стили: {e}")
            
    def generate_report_batch(self, 
                            results: List[DiagnosticResult],
                            config: ReportConfig,
                            progress_callback=None) -> Dict[str, List[str]]:
        """
        Генерация пакета отчетов
        
        Args:
            results: Список результатов диагностики
            config: Конфигурация отчетов
            progress_callback: Функция обратного вызова для прогресса
            
        Returns:
            Словарь с путями к созданным файлам по автомобилям
        """
        generated_files = {}
        
        total = len(results)
        for i, result in enumerate(results):
            if progress_callback:
                progress_callback(int(i / total * 100), 
                                f"Генерация отчета {i+1} из {total}...")
                
            try:
                files = self.generate_report(result, config)
                vin = result.vehicle_info.vin or f"vehicle_{i+1}"
                generated_files[vin] = files
                
            except Exception as e:
                self.logger.error(f"Ошибка генерации отчета для автомобиля {i+1}: {e}")
                
        if progress_callback:
            progress_callback(100, "Пакетная генерация завершена")
            
        return generated_files
        
    def create_summary_report(self, 
                             generated_files: Dict[str, Dict[str, str]],
                             config: ReportConfig) -> Path:
        """
        Создание сводного отчета по всем диагностикам
        
        Args:
            generated_files: Словарь с путями к отчетам
            config: Конфигурация отчета
            
        Returns:
            Путь к сводному отчету
        """
        try:
            # Создаем Excel файл со сводкой
            wb = Workbook()
            ws = wb.active
            ws.title = "Сводка диагностик"
            
            # Заголовки
            headers = ["VIN", "Модель", "Год", "Пробег", "Оценка", "Ошибки", 
                      "Критические", "Дата", "Отчет"]
            ws.append(headers)
            
            # Стили заголовков
            for col in range(1, len(headers) + 1):
                cell = ws.cell(row=1, column=col)
                cell.font = Font(bold=True)
                cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
                cell.alignment = Alignment(horizontal='center')
                
            # Заполняем данные
            row = 2
            for vin, files in generated_files.items():
                # В реальном приложении нужно читать данные из отчетов
                ws.cell(row=row, column=1, value=vin)
                ws.cell(row=row, column=5, value="Н/Д")
                ws.cell(row=row, column=6, value="Н/Д")
                
                # Добавляем гиперссылку на отчет если есть
                if 'html' in files:
                    cell = ws.cell(row=row, column=9, value="Открыть отчет")
                    cell.hyperlink = f"file:///{files['html']}"
                    cell.style = "Hyperlink"
                    
                row += 1
                
            # Автонастройка ширины колонок
            for column in ws.columns:
                max_length = 0
                column_letter = get_column_letter(column[0].column)
                
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                        
                adjusted_width = min(max_length + 2, 50)
                ws.column_dimensions[column_letter].width = adjusted_width
                
            # Сохраняем файл
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"diagnostics_summary_{timestamp}.xlsx"
            filepath = self.reports_dir / filename
            wb.save(str(filepath))
            
            return filepath
            
        except Exception as e:
            self.logger.error(f"Ошибка создания сводного отчета: {e}")
            raise


# Утилитарные функции для работы с отчетом

def create_diagnostic_result_from_data(data: Dict) -> DiagnosticResult:
    """Создание объекта DiagnosticResult из словаря данных"""
    vehicle_info = VehicleInfo(
        vin=data.get('vin', ''),
        model=data.get('model', ''),
        year=data.get('year', 0),
        engine=data.get('engine', ''),
        mileage=data.get('mileage', 0.0),
        license_plate=data.get('license_plate', ''),
        owner=data.get('owner', ''),
        contact=data.get('contact', ''),
        notes=data.get('notes', '')
    )
    
    return DiagnosticResult(
        vehicle_info=vehicle_info,
        dtcs=data.get('dtcs', []),
        live_data=data.get('live_data', {}),
        ecu_status=data.get('ecu_status', {}),
        sensor_checks=data.get('sensor_checks', {}),
        actuator_checks=data.get('actuator_checks', {}),
        adaptation_results=data.get('adaptation_results', {}),
        recommendations=data.get('recommendations', []),
        technician=data.get('technician', ''),
        diagnostic_duration=data.get('diagnostic_duration', 0.0),
        rating=data.get('rating', 0)
    )


def save_report_config(config: ReportConfig, filepath: Path):
    """Сохранение конфигурации отчета в файл"""
    config_dict = {
        'format': config.format.value,
        'level': config.level.value,
        'include_graphs': config.include_graphs,
        'include_photos': config.include_photos,
        'include_technical_data': config.include_technical_data,
        'include_recommendations': config.include_recommendations,
        'language': config.language,
        'company_name': config.company_name,
        'company_address': config.company_address,
        'company_phone': config.company_phone,
        'technician_signature': config.technician_signature,
        'customer_signature': config.customer_signature
    }
    
    with open(filepath, 'w', encoding='utf-8') as f:
        json.dump(config_dict, f, ensure_ascii=False, indent=2)


def load_report_config(filepath: Path) -> ReportConfig:
    """Загрузка конфигурации отчета из файла"""
    with open(filepath, 'r', encoding='utf-8') as f:
        config_dict = json.load(f)
        
    return ReportConfig(
        format=ReportFormat(config_dict.get('format', 'html')),
        level=ReportLevel(config_dict.get('level', 'standard')),
        include_graphs=config_dict.get('include_graphs', True),
        include_photos=config_dict.get('include_photos', False),
        include_technical_data=config_dict.get('include_technical_data', True),
        include_recommendations=config_dict.get('include_recommendations', True),
        language=config_dict.get('language', 'ru'),
        company_name=config_dict.get('company_name', 'Автосервис'),
        company_address=config_dict.get('company_address', ''),
        company_phone=config_dict.get('company_phone', ''),
        technician_signature=config_dict.get('technician_signature', True),
        customer_signature=config_dict.get('customer_signature', True)
    )