"""
Панель отчетов - генерация, сохранение и управление диагностическими отчетами
"""

import os
import json
from datetime import datetime
from PyQt5.QtWidgets import (QWidget, QVBoxLayout, QHBoxLayout, QGridLayout,
                             QPushButton, QLabel, QComboBox, QTextEdit,
                             QTableWidget, QTableWidgetItem, QHeaderView,
                             QGroupBox, QCheckBox, QSpinBox, QDoubleSpinBox,
                             QFileDialog, QMessageBox, QProgressBar,
                             QTabWidget, QSplitter, QFrame, QTreeWidget,
                             QTreeWidgetItem, QTreeView, QAbstractItemView,
                             QLineEdit, QDateEdit, QFormLayout, QListWidget,
                             QListWidgetItem, QDialog, QDialogButtonBox)
from PyQt5.QtCore import Qt, QTimer, QDate, pyqtSignal, QThread, pyqtSlot
from PyQt5.QtGui import QFont, QColor, QBrush, QIcon, QTextDocument, QPixmap
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm, inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import matplotlib.pyplot as plt
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.figure import Figure
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.drawing.image import Image as XLImage
import jinja2
import base64
import tempfile
import webbrowser
from io import BytesIO

from utils.logger import setup_logger
from config_manager import ConfigManager


class ReportGeneratorThread(QThread):
    """Поток для генерации отчетов"""
    progress = pyqtSignal(int)
    status = pyqtSignal(str)
    finished = pyqtSignal(str, bool)  # путь, успех
    error = pyqtSignal(str)
    
    def __init__(self, report_data, report_type, template_name, output_path):
        super().__init__()
        self.report_data = report_data
        self.report_type = report_type
        self.template_name = template_name
        self.output_path = output_path
        self.logger = setup_logger()
        
    def run(self):
        try:
            if self.report_type == "pdf":
                self._generate_pdf_report()
            elif self.report_type == "docx":
                self._generate_docx_report()
            elif self.report_type == "excel":
                self._generate_excel_report()
            elif self.report_type == "html":
                self._generate_html_report()
            else:
                self.error.emit(f"Неизвестный тип отчета: {self.report_type}")
                return
                
            self.finished.emit(self.output_path, True)
            
        except Exception as e:
            self.logger.error(f"Ошибка генерации отчета: {e}")
            self.error.emit(str(e))
            self.finished.emit("", False)
            
    def _generate_pdf_report(self):
        """Генерация PDF отчета"""
        self.status.emit("Создание PDF документа...")
        self.progress.emit(10)
        
        # Создаем документ
        doc = SimpleDocTemplate(
            self.output_path,
            pagesize=A4,
            rightMargin=20*mm,
            leftMargin=20*mm,
            topMargin=20*mm,
            bottomMargin=20*mm
        )
        
        # Регистрация шрифтов (если нужно)
        try:
            pdfmetrics.registerFont(TTFont('DejaVu', 'DejaVuSans.ttf'))
        except:
            pass
            
        styles = getSampleStyleSheet()
        
        # Создаем стиль заголовка
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=12,
            alignment=1  # центрирование
        )
        
        # Создаем стиль подзаголовка
        subtitle_style = ParagraphStyle(
            'CustomSubtitle',
            parent=styles['Heading2'],
            fontSize=12,
            spaceAfter=6,
            textColor=colors.grey
        )
        
        # Содержимое документа
        story = []
        
        # Заголовок
        title = Paragraph("ДИАГНОСТИЧЕСКИЙ ОТЧЕТ", title_style)
        story.append(title)
        
        subtitle = Paragraph(f"Автомобиль: {self.report_data.get('vehicle_model', 'N/A')}", subtitle_style)
        story.append(subtitle)
        
        story.append(Spacer(1, 10*mm))
        self.progress.emit(20)
        
        # Информация о диагностике
        info_data = [
            ["Дата диагностики:", self.report_data.get('timestamp', 'N/A')],
            ["Модель автомобиля:", self.report_data.get('vehicle_model', 'N/A')],
            ["VIN:", self.report_data.get('vin', 'N/A')],
            ["Пробег:", f"{self.report_data.get('mileage', 'N/A')} км"],
            ["Статус диагностики:", self.report_data.get('diagnostic_status', 'N/A')],
        ]
        
        info_table = Table(info_data, colWidths=[60*mm, 80*mm])
        info_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ('PADDING', (0, 0), (-1, -1), 6),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        
        story.append(info_table)
        story.append(Spacer(1, 10*mm))
        self.progress.emit(30)
        
        # Секция с ошибками
        if 'dtcs' in self.report_data and self.report_data['dtcs']:
            self.status.emit("Добавление информации об ошибках...")
            
            dtc_title = Paragraph("КОДЫ НЕИСПРАВНОСТЕЙ (DTC)", styles['Heading2'])
            story.append(dtc_title)
            
            # Создаем таблицу ошибок
            dtc_headers = ["ECU", "Код", "Описание", "Статус"]
            dtc_rows = []
            
            for ecu, dtc_list in self.report_data['dtcs'].items():
                if dtc_list:
                    for dtc in dtc_list:
                        dtc_rows.append([
                            ecu,
                            dtc.get('code', 'N/A'),
                            dtc.get('description', 'N/A'),
                            dtc.get('status', 'Активная')
                        ])
            
            if dtc_rows:
                dtc_table_data = [dtc_headers] + dtc_rows
                dtc_table = Table(dtc_table_data, colWidths=[30*mm, 30*mm, 80*mm, 30*mm])
                dtc_table.setStyle(TableStyle([
                    ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 0), (-1, -1), 9),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('BACKGROUND', (0, 0), (-1, 0), colors.darkgrey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('PADDING', (0, 0), (-1, -1), 4),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.whitesmoke]),
                ]))
                
                story.append(dtc_table)
                
            story.append(Spacer(1, 10*mm))
            
        self.progress.emit(50)
        
        # Секция с параметрами
        if 'live_data' in self.report_data and self.report_data['live_data']:
            self.status.emit("Добавление параметров...")
            
            params_title = Paragraph("ТЕКУЩИЕ ПАРАМЕТРЫ", styles['Heading2'])
            story.append(params_title)
            
            # Группируем параметры по категориям
            param_categories = {}
            for param_name, param_data in self.report_data['live_data'].items():
                category = param_name.split('_')[0] if '_' in param_name else 'Общие'
                if category not in param_categories:
                    param_categories[category] = []
                    
                param_categories[category].append((
                    param_name,
                    param_data.get('value', 'N/A'),
                    param_data.get('unit', ''),
                    param_data.get('min', ''),
                    param_data.get('max', ''),
                    param_data.get('status', 'OK')
                ))
            
            # Создаем таблицы для каждой категории
            for category, params in param_categories.items():
                cat_title = Paragraph(category.upper(), styles['Heading3'])
                story.append(cat_title)
                
                param_headers = ["Параметр", "Значение", "Ед. изм.", "Мин.", "Макс.", "Статус"]
                param_rows = []
                
                for param in params:
                    status_color = colors.green if param[5] == 'OK' else colors.red
                    param_rows.append([
                        param[0],
                        str(param[1]),
                        param[2],
                        str(param[3]) if param[3] != '' else 'N/A',
                        str(param[4]) if param[4] != '' else 'N/A',
                        param[5]
                    ])
                
                param_table_data = [param_headers] + param_rows
                param_table = Table(param_table_data, colWidths=[50*mm, 30*mm, 20*mm, 20*mm, 20*mm, 20*mm])
                
                # Определяем цвета строк в зависимости от статуса
                row_colors = []
                for i in range(1, len(param_table_data)):
                    status = param_table_data[i][5]
                    if status == 'OK':
                        row_colors.append(colors.white)
                    else:
                        row_colors.append(colors.lightpink)
                
                param_table.setStyle(TableStyle([
                    ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 0), (-1, -1), 9),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('BACKGROUND', (0, 0), (-1, 0), colors.lightblue),
                    ('PADDING', (0, 0), (-1, -1), 4),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), row_colors),
                ]))
                
                story.append(param_table)
                story.append(Spacer(1, 5*mm))
                
        self.progress.emit(70)
        
        # Секция с рекомендациями
        if 'recommendations' in self.report_data and self.report_data['recommendations']:
            self.status.emit("Добавление рекомендаций...")
            
            rec_title = Paragraph("РЕКОМЕНДАЦИИ", styles['Heading2'])
            story.append(rec_title)
            
            for i, rec in enumerate(self.report_data['recommendations'], 1):
                rec_text = f"{i}. {rec.get('text', '')} [Приоритет: {rec.get('priority', 'Средний')}]"
                rec_para = Paragraph(rec_text, styles['Normal'])
                story.append(rec_para)
                
            story.append(Spacer(1, 10*mm))
            
        self.progress.emit(80)
        
        # Секция с графиками
        if 'charts' in self.report_data and self.report_data['charts']:
            self.status.emit("Добавление графиков...")
            
            charts_title = Paragraph("ГРАФИКИ ПАРАМЕТРОВ", styles['Heading2'])
            story.append(charts_title)
            
            for chart_name, chart_data in self.report_data['charts'].items():
                try:
                    # Создаем временный файл с изображением графика
                    temp_chart = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
                    self._create_chart_image(chart_name, chart_data, temp_chart.name)
                    temp_chart.close()
                    
                    # Добавляем изображение в отчет
                    chart_img = Image(temp_chart.name, width=150*mm, height=80*mm)
                    story.append(chart_img)
                    story.append(Spacer(1, 5*mm))
                    
                    # Удаляем временный файл
                    os.unlink(temp_chart.name)
                    
                except Exception as e:
                    self.logger.warning(f"Не удалось добавить график {chart_name}: {e}")
                    
        self.progress.emit(90)
        
        # Заключительная информация
        summary_title = Paragraph("ИТОГИ ДИАГНОСТИКИ", styles['Heading2'])
        story.append(summary_title)
        
        summary_text = f"""
        Всего проверено систем: {self.report_data.get('systems_checked', 0)}
        Найдено ошибок: {self.report_data.get('errors_found', 0)}
        Критических ошибок: {self.report_data.get('critical_errors', 0)}
        Рекомендаций: {len(self.report_data.get('recommendations', []))}
        Общее состояние: {self.report_data.get('overall_status', 'N/A')}
        """
        
        summary_para = Paragraph(summary_text, styles['Normal'])
        story.append(summary_para)
        story.append(Spacer(1, 10*mm))
        
        # Подпись
        signature = Paragraph(f"Отчет сгенерирован: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}", 
                             ParagraphStyle('Signature', parent=styles['Normal'], fontSize=8, textColor=colors.grey))
        story.append(signature)
        
        self.status.emit("Сохранение PDF документа...")
        self.progress.emit(95)
        
        # Сохраняем документ
        doc.build(story)
        self.progress.emit(100)
        
    def _create_chart_image(self, chart_name, chart_data, output_path):
        """Создание изображения графика"""
        fig, ax = plt.subplots(figsize=(10, 6))
        
        if 'x_data' in chart_data and 'y_data' in chart_data:
            x = chart_data['x_data']
            y = chart_data['y_data']
            
            ax.plot(x, y, marker='o', linewidth=2, markersize=4)
            
            # Добавляем сетку и заголовки
            ax.grid(True, alpha=0.3)
            ax.set_title(chart_name, fontsize=14)
            ax.set_xlabel(chart_data.get('x_label', 'Время'), fontsize=12)
            ax.set_ylabel(chart_data.get('y_label', 'Значение'), fontsize=12)
            
            # Настройка внешнего вида
            plt.xticks(rotation=45)
            plt.tight_layout()
            
        plt.savefig(output_path, dpi=150, bbox_inches='tight')
        plt.close(fig)
        
    def _generate_docx_report(self):
        """Генерация Word отчета"""
        self.status.emit("Создание Word документа...")
        self.progress.emit(10)
        
        doc = Document()
        
        # Настройка стилей
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        # Заголовок
        title = doc.add_heading('ДИАГНОСТИЧЕСКИЙ ОТЧЕТ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Подзаголовок
        subtitle = doc.add_paragraph()
        subtitle_run = subtitle.add_run(f'Автомобиль: {self.report_data.get("vehicle_model", "N/A")}')
        subtitle_run.font.size = Pt(12)
        subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        self.progress.emit(20)
        
        # Информация о диагностике
        info_title = doc.add_heading('Информация о диагностике', 2)
        
        # Создаем таблицу с информацией
        info_table = doc.add_table(rows=5, cols=2)
        info_table.style = 'Light Grid'
        info_table.autofit = False
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        info_rows = [
            ["Дата диагностики:", self.report_data.get('timestamp', 'N/A')],
            ["Модель автомобиля:", self.report_data.get('vehicle_model', 'N/A')],
            ["VIN:", self.report_data.get('vin', 'N/A')],
            ["Пробег:", f"{self.report_data.get('mileage', 'N/A')} км"],
            ["Статус диагностики:", self.report_data.get('diagnostic_status', 'N/A')],
        ]
        
        for i, row in enumerate(info_rows):
            cells = info_table.rows[i].cells
            cells[0].text = row[0]
            cells[1].text = row[1]
            
            # Жирный шрифт для заголовков
            cells[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        self.progress.emit(30)
        
        # Секция с ошибками
        if 'dtcs' in self.report_data and self.report_data['dtcs']:
            self.status.emit("Добавление информации об ошибках...")
            
            dtc_title = doc.add_heading('Коды неисправностей (DTC)', 2)
            
            for ecu, dtc_list in self.report_data['dtcs'].items():
                if dtc_list:
                    ecu_title = doc.add_heading(f'Модуль: {ecu}', 3)
                    
                    # Создаем таблицу ошибок для каждого ECU
                    dtc_table = doc.add_table(rows=len(dtc_list) + 1, cols=4)
                    dtc_table.style = 'Light Grid'
                    
                    # Заголовки таблицы
                    headers = dtc_table.rows[0].cells
                    headers[0].text = 'Код'
                    headers[1].text = 'Описание'
                    headers[2].text = 'Статус'
                    headers[3].text = 'Дата'
                    
                    # Делаем заголовки жирными
                    for header in headers:
                        header.paragraphs[0].runs[0].font.bold = True
                    
                    # Заполняем данные
                    for i, dtc in enumerate(dtc_list, 1):
                        cells = dtc_table.rows[i].cells
                        cells[0].text = dtc.get('code', 'N/A')
                        cells[1].text = dtc.get('description', 'N/A')
                        cells[2].text = dtc.get('status', 'Активная')
                        cells[3].text = dtc.get('date', 'N/A')
                        
                        # Подсветка критических ошибок
                        if dtc.get('critical', False):
                            for cell in cells:
                                shading = cell.paragraphs[0].runs[0].font
                                shading.color.rgb = RGBColor(255, 0, 0)
                    
                    doc.add_paragraph()
            
        self.progress.emit(50)
        
        # Секция с параметрами
        if 'live_data' in self.report_data and self.report_data['live_data']:
            self.status.emit("Добавление параметров...")
            
            params_title = doc.add_heading('Текущие параметры', 2)
            
            # Группируем параметры
            param_groups = {}
            for param_name, param_data in self.report_data['live_data'].items():
                group = param_name.split('_')[0] if '_' in param_name else 'Общие'
                if group not in param_groups:
                    param_groups[group] = []
                    
                param_groups[group].append({
                    'name': param_name,
                    'value': param_data.get('value', 'N/A'),
                    'unit': param_data.get('unit', ''),
                    'min': param_data.get('min', ''),
                    'max': param_data.get('max', ''),
                    'status': param_data.get('status', 'OK')
                })
            
            # Создаем таблицы для каждой группы
            for group_name, params in param_groups.items():
                group_title = doc.add_heading(group_name, 3)
                
                param_table = doc.add_table(rows=len(params) + 1, cols=6)
                param_table.style = 'Light Grid'
                
                # Заголовки
                headers = param_table.rows[0].cells
                headers[0].text = 'Параметр'
                headers[1].text = 'Значение'
                headers[2].text = 'Ед. изм.'
                headers[3].text = 'Мин.'
                headers[4].text = 'Макс.'
                headers[5].text = 'Статус'
                
                for header in headers:
                    header.paragraphs[0].runs[0].font.bold = True
                
                # Данные
                for i, param in enumerate(params, 1):
                    cells = param_table.rows[i].cells
                    cells[0].text = param['name']
                    cells[1].text = str(param['value'])
                    cells[2].text = param['unit']
                    cells[3].text = str(param['min']) if param['min'] != '' else 'N/A'
                    cells[4].text = str(param['max']) if param['max'] != '' else 'N/A'
                    cells[5].text = param['status']
                    
                    # Подсветка проблемных параметров
                    if param['status'] != 'OK':
                        for cell in cells:
                            shading = cell.paragraphs[0].runs[0].font
                            shading.color.rgb = RGBColor(255, 0, 0)
                
                doc.add_paragraph()
                
        self.progress.emit(70)
        
        # Рекомендации
        if 'recommendations' in self.report_data and self.report_data['recommendations']:
            self.status.emit("Добавление рекомендаций...")
            
            rec_title = doc.add_heading('Рекомендации', 2)
            
            for i, rec in enumerate(self.report_data['recommendations'], 1):
                rec_text = doc.add_paragraph()
                rec_text.add_run(f'{i}. {rec.get("text", "")} ').bold = True
                
                priority_text = f'[Приоритет: {rec.get("priority", "Средний")}]'
                priority_run = rec_text.add_run(priority_text)
                
                # Цвет приоритета
                priority = rec.get('priority', 'Средний')
                if priority == 'Высокий':
                    priority_run.font.color.rgb = RGBColor(255, 0, 0)
                elif priority == 'Средний':
                    priority_run.font.color.rgb = RGBColor(255, 165, 0)
                else:
                    priority_run.font.color.rgb = RGBColor(0, 128, 0)
                
                if 'details' in rec and rec['details']:
                    details = doc.add_paragraph(rec['details'])
                    details.style = 'Intense Quote'
            
            doc.add_paragraph()
            
        self.progress.emit(80)
        
        # Итоги
        summary_title = doc.add_heading('Итоги диагностики', 2)
        
        summary_data = [
            ["Всего проверено систем:", str(self.report_data.get('systems_checked', 0))],
            ["Найдено ошибок:", str(self.report_data.get('errors_found', 0))],
            ["Критических ошибок:", str(self.report_data.get('critical_errors', 0))],
            ["Рекомендаций:", str(len(self.report_data.get('recommendations', [])))],
            ["Общее состояние:", self.report_data.get('overall_status', 'N/A')],
        ]
        
        summary_table = doc.add_table(rows=5, cols=2)
        summary_table.style = 'Light Grid'
        
        for i, row in enumerate(summary_data):
            cells = summary_table.rows[i].cells
            cells[0].text = row[0]
            cells[1].text = row[1]
            cells[0].paragraphs[0].runs[0].font.bold = True
            
            if i == 4:  # Последняя строка - общее состояние
                color = RGBColor(0, 128, 0) if row[1] == 'Хорошее' else RGBColor(255, 0, 0)
                cells[1].paragraphs[0].runs[0].font.color.rgb = color
        
        doc.add_paragraph()
        
        # Подпись
        signature = doc.add_paragraph()
        signature.add_run(f'Отчет сгенерирован: {datetime.now().strftime("%d.%m.%Y %H:%M:%S")}')
        signature.runs[0].font.size = Pt(9)
        signature.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        self.status.emit("Сохранение Word документа...")
        self.progress.emit(90)
        
        # Сохраняем документ
        doc.save(self.output_path)
        self.progress.emit(100)
        
    def _generate_excel_report(self):
        """Генерация Excel отчета"""
        self.status.emit("Создание Excel документа...")
        self.progress.emit(10)
        
        wb = Workbook()
        ws = wb.active
        ws.title = "Диагностический отчет"
        
        # Стили
        header_font = Font(name='Calibri', size=14, bold=True, color='FFFFFF')
        header_fill = PatternFill(start_color='366092', end_color='366092', fill_type='solid')
        header_alignment = Alignment(horizontal='center', vertical='center')
        
        subheader_font = Font(name='Calibri', size=12, bold=True, color='000000')
        subheader_fill = PatternFill(start_color='D9E1F2', end_color='D9E1F2', fill_type='solid')
        
        normal_font = Font(name='Calibri', size=11)
        border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # Заголовок
        ws.merge_cells('A1:F1')
        title_cell = ws['A1']
        title_cell.value = "ДИАГНОСТИЧЕСКИЙ ОТЧЕТ"
        title_cell.font = header_font
        title_cell.fill = header_fill
        title_cell.alignment = header_alignment
        
        ws.merge_cells('A2:F2')
        subtitle_cell = ws['A2']
        subtitle_cell.value = f"Автомобиль: {self.report_data.get('vehicle_model', 'N/A')}"
        subtitle_cell.font = subheader_font
        subtitle_cell.alignment = Alignment(horizontal='center')
        
        self.progress.emit(20)
        
        # Информация о диагностике
        current_row = 4
        
        ws.merge_cells(f'A{current_row}:F{current_row}')
        info_title = ws.cell(row=current_row, column=1)
        info_title.value = "Информация о диагностике"
        info_title.font = subheader_font
        info_title.fill = subheader_fill
        current_row += 1
        
        info_data = [
            ["Дата диагностики:", self.report_data.get('timestamp', 'N/A')],
            ["Модель автомобиля:", self.report_data.get('vehicle_model', 'N/A')],
            ["VIN:", self.report_data.get('vin', 'N/A')],
            ["Пробег:", f"{self.report_data.get('mileage', 'N/A')} км"],
            ["Статус диагностики:", self.report_data.get('diagnostic_status', 'N/A')],
        ]
        
        for i, row_data in enumerate(info_data):
            for j, cell_data in enumerate(row_data):
                cell = ws.cell(row=current_row + i, column=j + 1)
                cell.value = cell_data
                cell.border = border
                cell.font = normal_font
                
                if j == 0:  # Заголовки
                    cell.font = Font(name='Calibri', size=11, bold=True)
        
        current_row += len(info_data) + 2
        self.progress.emit(30)
        
        # Ошибки
        if 'dtcs' in self.report_data and self.report_data['dtcs']:
            self.status.emit("Добавление информации об ошибках...")
            
            ws.merge_cells(f'A{current_row}:F{current_row}')
            dtc_title = ws.cell(row=current_row, column=1)
            dtc_title.value = "Коды неисправностей (DTC)"
            dtc_title.font = subheader_font
            dtc_title.fill = subheader_fill
            current_row += 1
            
            dtc_headers = ["ECU", "Код", "Описание", "Статус", "Дата", "Критичность"]
            
            for i, header in enumerate(dtc_headers):
                cell = ws.cell(row=current_row, column=i + 1)
                cell.value = header
                cell.font = Font(name='Calibri', size=11, bold=True)
                cell.fill = PatternFill(start_color='E2EFDA', end_color='E2EFDA', fill_type='solid')
                cell.border = border
                cell.alignment = Alignment(horizontal='center')
            
            current_row += 1
            
            row_index = current_row
            for ecu, dtc_list in self.report_data['dtcs'].items():
                if dtc_list:
                    for dtc in dtc_list:
                        dtc_row = [
                            ecu,
                            dtc.get('code', 'N/A'),
                            dtc.get('description', 'N/A'),
                            dtc.get('status', 'Активная'),
                            dtc.get('date', 'N/A'),
                            dtc.get('critical', 'Нет')
                        ]
                        
                        for i, value in enumerate(dtc_row):
                            cell = ws.cell(row=row_index, column=i + 1)
                            cell.value = value
                            cell.border = border
                            cell.font = normal_font
                            
                            if dtc.get('critical', False):
                                cell.font = Font(name='Calibri', size=11, bold=True, color='FF0000')
                        
                        row_index += 1
            
            current_row = row_index + 2
            
        self.progress.emit(50)
        
        # Параметры
        if 'live_data' in self.report_data and self.report_data['live_data']:
            self.status.emit("Добавление параметров...")
            
            ws.merge_cells(f'A{current_row}:F{current_row}')
            params_title = ws.cell(row=current_row, column=1)
            params_title.value = "Текущие параметры"
            params_title.font = subheader_font
            params_title.fill = subheader_fill
            current_row += 1
            
            param_headers = ["Параметр", "Значение", "Ед. изм.", "Мин.", "Макс.", "Статус"]
            
            for i, header in enumerate(param_headers):
                cell = ws.cell(row=current_row, column=i + 1)
                cell.value = header
                cell.font = Font(name='Calibri', size=11, bold=True)
                cell.fill = PatternFill(start_color='FFF2CC', end_color='FFF2CC', fill_type='solid')
                cell.border = border
                cell.alignment = Alignment(horizontal='center')
            
            current_row += 1
            
            row_index = current_row
            for param_name, param_data in self.report_data['live_data'].items():
                param_row = [
                    param_name,
                    param_data.get('value', 'N/A'),
                    param_data.get('unit', ''),
                    param_data.get('min', 'N/A'),
                    param_data.get('max', 'N/A'),
                    param_data.get('status', 'OK')
                ]
                
                for i, value in enumerate(param_row):
                    cell = ws.cell(row=row_index, column=i + 1)
                    cell.value = value
                    cell.border = border
                    cell.font = normal_font
                    
                    if param_data.get('status', 'OK') != 'OK':
                        cell.font = Font(name='Calibri', size=11, bold=True, color='FF0000')
                        cell.fill = PatternFill(start_color='FFC7CE', end_color='FFC7CE', fill_type='solid')
                
                row_index += 1
            
            # Автоподбор ширины столбцов
            for column in ws.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                ws.column_dimensions[column_letter].width = adjusted_width
            
            current_row = row_index + 2
            
        self.progress.emit(70)
        
        # Рекомендации
        if 'recommendations' in self.report_data and self.report_data['recommendations']:
            self.status.emit("Добавление рекомендаций...")
            
            ws.merge_cells(f'A{current_row}:F{current_row}')
            rec_title = ws.cell(row=current_row, column=1)
            rec_title.value = "Рекомендации"
            rec_title.font = subheader_font
            rec_title.fill = subheader_fill
            current_row += 1
            
            for i, rec in enumerate(self.report_data['recommendations'], 1):
                rec_text = f"{i}. {rec.get('text', '')}"
                ws.merge_cells(f'A{current_row}:F{current_row}')
                rec_cell = ws.cell(row=current_row, column=1)
                rec_cell.value = rec_text
                rec_cell.border = border
                rec_cell.font = normal_font
                
                priority = rec.get('priority', 'Средний')
                if priority == 'Высокий':
                    rec_cell.font = Font(name='Calibri', size=11, bold=True, color='FF0000')
                    rec_cell.fill = PatternFill(start_color='FFC7CE', end_color='FFC7CE', fill_type='solid')
                elif priority == 'Средний':
                    rec_cell.font = Font(name='Calibri', size=11, bold=True, color='FF9900')
                    rec_cell.fill = PatternFill(start_color='FFE699', end_color='FFE699', fill_type='solid')
                
                current_row += 1
                
                if 'details' in rec and rec['details']:
                    ws.merge_cells(f'B{current_row}:F{current_row}')
                    details_cell = ws.cell(row=current_row, column=2)
                    details_cell.value = rec['details']
                    details_cell.font = Font(name='Calibri', size=10, italic=True)
                    details_cell.alignment = Alignment(wrap_text=True)
                    current_row += 1
                
                current_row += 1
            
        self.progress.emit(85)
        
        # Сводка
        ws.merge_cells(f'A{current_row}:F{current_row}')
        summary_title = ws.cell(row=current_row, column=1)
        summary_title.value = "Сводка диагностики"
        summary_title.font = subheader_font
        summary_title.fill = subheader_fill
        current_row += 1
        
        summary_data = [
            ["Всего проверено систем:", self.report_data.get('systems_checked', 0)],
            ["Найдено ошибок:", self.report_data.get('errors_found', 0)],
            ["Критических ошибок:", self.report_data.get('critical_errors', 0)],
            ["Рекомендаций:", len(self.report_data.get('recommendations', []))],
            ["Общее состояние:", self.report_data.get('overall_status', 'N/A')],
        ]
        
        for i, row_data in enumerate(summary_data):
            for j, cell_data in enumerate(row_data):
                cell = ws.cell(row=current_row + i, column=j + 1)
                cell.value = cell_data
                cell.border = border
                cell.font = normal_font
                
                if j == 0:
                    cell.font = Font(name='Calibri', size=11, bold=True)
                elif i == 4:  # Общее состояние
                    if cell_data == 'Хорошее':
                        cell.font = Font(name='Calibri', size=11, bold=True, color='00B050')
                    else:
                        cell.font = Font(name='Calibri', size=11, bold=True, color='FF0000')
        
        current_row += len(summary_data) + 2
        
        # Подпись
        signature = f"Отчет сгенерирован: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}"
        ws.merge_cells(f'A{current_row}:F{current_row}')
        signature_cell = ws.cell(row=current_row, column=1)
        signature_cell.value = signature
        signature_cell.font = Font(name='Calibri', size=9, color='808080')
        signature_cell.alignment = Alignment(horizontal='right')
        
        self.status.emit("Сохранение Excel документа...")
        self.progress.emit(95)
        
        # Сохраняем файл
        wb.save(self.output_path)
        self.progress.emit(100)
        
    def _generate_html_report(self):
        """Генерация HTML отчета"""
        self.status.emit("Создание HTML документа...")
        self.progress.emit(10)
        
        # Загружаем шаблон Jinja2
        template_str = self._get_html_template()
        template = jinja2.Template(template_str)
        
        # Подготавливаем данные для шаблона
        context = {
            'report_data': self.report_data,
            'generation_date': datetime.now().strftime('%d.%m.%Y %H:%M:%S'),
            'styles': self._get_html_styles(),
            'has_errors': bool(self.report_data.get('dtcs')),
            'has_params': bool(self.report_data.get('live_data')),
            'has_recommendations': bool(self.report_data.get('recommendations')),
        }
        
        self.progress.emit(30)
        
        # Рендерим HTML
        html_content = template.render(**context)
        
        self.progress.emit(60)
        
        # Сохраняем HTML файл
        with open(self.output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
            
        self.progress.emit(80)
        
        # Сохраняем связанные ресурсы (CSS, JS, изображения)
        self._save_html_resources(os.path.dirname(self.output_path))
        
        self.progress.emit(100)
        
    def _get_html_template(self):
        """Возвращает HTML шаблон"""
        return """
<!DOCTYPE html>
<html lang="ru">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Диагностический отчет - {{ report_data.vehicle_model }}</title>
    <style>
        {{ styles }}
    </style>
</head>
<body>
    <div class="container">
        <!-- Заголовок -->
        <header class="report-header">
            <h1>ДИАГНОСТИЧЕСКИЙ ОТЧЕТ</h1>
            <h2>Автомобиль: {{ report_data.vehicle_model }}</h2>
        </header>
        
        <!-- Информация о диагностике -->
        <section class="section">
            <h3>Информация о диагностике</h3>
            <div class="info-grid">
                <div class="info-item">
                    <span class="info-label">Дата диагностики:</span>
                    <span class="info-value">{{ report_data.timestamp }}</span>
                </div>
                <div class="info-item">
                    <span class="info-label">Модель автомобиля:</span>
                    <span class="info-value">{{ report_data.vehicle_model }}</span>
                </div>
                <div class="info-item">
                    <span class="info-label">VIN:</span>
                    <span class="info-value">{{ report_data.vin }}</span>
                </div>
                <div class="info-item">
                    <span class="info-label">Пробег:</span>
                    <span class="info-value">{{ report_data.mileage }} км</span>
                </div>
                <div class="info-item">
                    <span class="info-label">Статус диагностики:</span>
                    <span class="info-value {{ report_data.diagnostic_status|lower }}">{{ report_data.diagnostic_status }}</span>
                </div>
            </div>
        </section>
        
        {% if has_errors %}
        <!-- Коды ошибок -->
        <section class="section">
            <h3>Коды неисправностей (DTC)</h3>
            {% for ecu, dtc_list in report_data.dtcs.items() %}
                {% if dtc_list %}
                <div class="ecu-section">
                    <h4>Модуль: {{ ecu }}</h4>
                    <table class="data-table">
                        <thead>
                            <tr>
                                <th>Код</th>
                                <th>Описание</th>
                                <th>Статус</th>
                                <th>Дата</th>
                                <th>Критичность</th>
                            </tr>
                        </thead>
                        <tbody>
                            {% for dtc in dtc_list %}
                            <tr class="{{ 'critical' if dtc.critical else '' }}">
                                <td>{{ dtc.code }}</td>
                                <td>{{ dtc.description }}</td>
                                <td>{{ dtc.status }}</td>
                                <td>{{ dtc.date }}</td>
                                <td>{{ 'Критическая' if dtc.critical else 'Некритическая' }}</td>
                            </tr>
                            {% endfor %}
                        </tbody>
                    </table>
                </div>
                {% endif %}
            {% endfor %}
        </section>
        {% endif %}
        
        {% if has_params %}
        <!-- Текущие параметры -->
        <section class="section">
            <h3>Текущие параметры</h3>
            {% set param_groups = {} %}
            {% for param_name, param_data in report_data.live_data.items() %}
                {% set group = param_name.split('_')[0] if '_' in param_name else 'Общие' %}
                {% if group not in param_groups %}
                    {% set _ = param_groups.update({group: []}) %}
                {% endif %}
                {% set _ = param_groups[group].append({'name': param_name, 'data': param_data}) %}
            {% endfor %}
            
            {% for group_name, params in param_groups.items() %}
            <div class="param-group">
                <h4>{{ group_name }}</h4>
                <table class="data-table">
                    <thead>
                        <tr>
                            <th>Параметр</th>
                            <th>Значение</th>
                            <th>Ед. изм.</th>
                            <th>Мин.</th>
                            <th>Макс.</th>
                            <th>Статус</th>
                        </tr>
                    </thead>
                    <tbody>
                        {% for param in params %}
                        <tr class="{{ 'warning' if param.data.status != 'OK' else '' }}">
                            <td>{{ param.name }}</td>
                            <td>{{ param.data.value }}</td>
                            <td>{{ param.data.unit }}</td>
                            <td>{{ param.data.min or 'N/A' }}</td>
                            <td>{{ param.data.max or 'N/A' }}</td>
                            <td class="status-{{ param.data.status|lower }}">{{ param.data.status }}</td>
                        </tr>
                        {% endfor %}
                    </tbody>
                </table>
            </div>
            {% endfor %}
        </section>
        {% endif %}
        
        {% if has_recommendations %}
        <!-- Рекомендации -->
        <section class="section">
            <h3>Рекомендации</h3>
            <div class="recommendations">
                {% for rec in report_data.recommendations %}
                <div class="recommendation priority-{{ rec.priority|lower }}">
                    <div class="rec-header">
                        <span class="rec-number">{{ loop.index }}.</span>
                        <span class="rec-text">{{ rec.text }}</span>
                        <span class="rec-priority">[Приоритет: {{ rec.priority }}]</span>
                    </div>
                    {% if rec.details %}
                    <div class="rec-details">{{ rec.details }}</div>
                    {% endif %}
                </div>
                {% endfor %}
            </div>
        </section>
        {% endif %}
        
        <!-- Сводка -->
        <section class="section summary">
            <h3>Сводка диагностики</h3>
            <div class="summary-grid">
                <div class="summary-item">
                    <span class="summary-label">Всего проверено систем:</span>
                    <span class="summary-value">{{ report_data.systems_checked }}</span>
                </div>
                <div class="summary-item">
                    <span class="summary-label">Найдено ошибок:</span>
                    <span class="summary-value">{{ report_data.errors_found }}</span>
                </div>
                <div class="summary-item">
                    <span class="summary-label">Критических ошибок:</span>
                    <span class="summary-value">{{ report_data.critical_errors }}</span>
                </div>
                <div class="summary-item">
                    <span class="summary-label">Рекомендаций:</span>
                    <span class="summary-value">{{ report_data.recommendations|length }}</span>
                </div>
                <div class="summary-item overall">
                    <span class="summary-label">Общее состояние:</span>
                    <span class="summary-value status-{{ report_data.overall_status|lower }}">{{ report_data.overall_status }}</span>
                </div>
            </div>
        </section>
        
        <!-- Подпись -->
        <footer class="report-footer">
            <p>Отчет сгенерирован: {{ generation_date }}</p>
        </footer>
    </div>
    
    <script>
        // Функции для интерактивности
        document.addEventListener('DOMContentLoaded', function() {
            // Добавляем обработчики для сворачивания секций
            const sectionHeaders = document.querySelectorAll('h3, h4');
            sectionHeaders.forEach(header => {
                header.addEventListener('click', function() {
                    const section = this.parentElement;
                    const content = section.querySelector('.info-grid, .data-table, .recommendations, .summary-grid');
                    if (content) {
                        content.classList.toggle('collapsed');
                    }
                });
            });
            
            // Подсветка строк таблиц при наведении
            const tableRows = document.querySelectorAll('.data-table tbody tr');
            tableRows.forEach(row => {
                row.addEventListener('mouseenter', function() {
                    this.style.backgroundColor = '#f5f5f5';
                });
                row.addEventListener('mouseleave', function() {
                    this.style.backgroundColor = '';
                });
            });
            
            // Кнопка печати (если нужна)
            const printButton = document.createElement('button');
            printButton.textContent = 'Печать отчета';
            printButton.className = 'print-button';
            printButton.onclick = () => window.print();
            document.querySelector('.report-footer').prepend(printButton);
        });
    </script>
</body>
</html>
        """
        
    def _get_html_styles(self):
        """Возвращает CSS стили для HTML отчета"""
        return """
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }
        
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            color: #333;
            background-color: #f8f9fa;
            padding: 20px;
        }
        
        .container {
            max-width: 1200px;
            margin: 0 auto;
            background: white;
            box-shadow: 0 0 20px rgba(0,0,0,0.1);
            border-radius: 10px;
            overflow: hidden;
        }
        
        .report-header {
            background: linear-gradient(135deg, #2c3e50, #4a6491);
            color: white;
            padding: 40px;
            text-align: center;
        }
        
        .report-header h1 {
            font-size: 32px;
            margin-bottom: 10px;
            font-weight: 700;
        }
        
        .report-header h2 {
            font-size: 20px;
            font-weight: 400;
            opacity: 0.9;
        }
        
        .section {
            padding: 30px;
            border-bottom: 1px solid #eaeaea;
        }
        
        .section:last-child {
            border-bottom: none;
        }
        
        .section h3 {
            font-size: 24px;
            color: #2c3e50;
            margin-bottom: 20px;
            padding-bottom: 10px;
            border-bottom: 2px solid #3498db;
            cursor: pointer;
            transition: color 0.3s;
        }
        
        .section h3:hover {
            color: #2980b9;
        }
        
        .section h4 {
            font-size: 18px;
            color: #34495e;
            margin: 15px 0;
            cursor: pointer;
        }
        
        .info-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(300px, 1fr));
            gap: 20px;
        }
        
        .info-item {
            display: flex;
            justify-content: space-between;
            padding: 15px;
            background: #f8f9fa;
            border-radius: 8px;
            border-left: 4px solid #3498db;
        }
        
        .info-label {
            font-weight: 600;
            color: #2c3e50;
        }
        
        .info-value {
            color: #7f8c8d;
            text-align: right;
        }
        
        .data-table {
            width: 100%;
            border-collapse: collapse;
            margin: 15px 0;
            font-size: 14px;
        }
        
        .data-table th {
            background: #34495e;
            color: white;
            padding: 12px;
            text-align: left;
            font-weight: 600;
        }
        
        .data-table td {
            padding: 10px 12px;
            border-bottom: 1px solid #eaeaea;
        }
        
        .data-table tbody tr:hover {
            background-color: #f5f5f5;
        }
        
        .data-table tbody tr.critical {
            background-color: #ffeaea;
        }
        
        .data-table tbody tr.critical:hover {
            background-color: #ffd6d6;
        }
        
        .data-table tbody tr.warning {
            background-color: #fff3cd;
        }
        
        .status-ok {
            color: #28a745;
            font-weight: 600;
        }
        
        .status-error, .status-critical {
            color: #dc3545;
            font-weight: 600;
        }
        
        .status-warning {
            color: #ffc107;
            font-weight: 600;
        }
        
        .param-group {
            margin-bottom: 30px;
        }
        
        .recommendations {
            display: flex;
            flex-direction: column;
            gap: 15px;
        }
        
        .recommendation {
            padding: 20px;
            border-radius: 8px;
            border-left: 4px solid #3498db;
        }
        
        .recommendation.priority-высокий {
            border-left-color: #dc3545;
            background: #ffeaea;
        }
        
        .recommendation.priority-средний {
            border-left-color: #ffc107;
            background: #fff3cd;
        }
        
        .recommendation.priority-низкий {
            border-left-color: #28a745;
            background: #d4edda;
        }
        
        .rec-header {
            display: flex;
            align-items: center;
            gap: 10px;
            margin-bottom: 10px;
        }
        
        .rec-number {
            font-weight: 700;
            color: #2c3e50;
        }
        
        .rec-text {
            flex: 1;
            font-weight: 600;
        }
        
        .rec-priority {
            font-style: italic;
            color: #6c757d;
        }
        
        .rec-details {
            padding-left: 30px;
            color: #6c757d;
            font-size: 14px;
            line-height: 1.5;
        }
        
        .summary {
            background: #f8f9fa;
        }
        
        .summary-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 20px;
        }
        
        .summary-item {
            display: flex;
            flex-direction: column;
            align-items: center;
            text-align: center;
            padding: 20px;
            background: white;
            border-radius: 8px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }
        
        .summary-item.overall {
            grid-column: 1 / -1;
            margin-top: 10px;
        }
        
        .summary-label {
            font-size: 14px;
            color: #6c757d;
            margin-bottom: 5px;
        }
        
        .summary-value {
            font-size: 24px;
            font-weight: 700;
            color: #2c3e50;
        }
        
        .report-footer {
            padding: 20px;
            text-align: center;
            color: #6c757d;
            font-size: 14px;
            border-top: 1px solid #eaeaea;
        }
        
        .print-button {
            background: #3498db;
            color: white;
            border: none;
            padding: 10px 20px;
            border-radius: 5px;
            cursor: pointer;
            font-size: 14px;
            transition: background 0.3s;
            margin-right: 10px;
        }
        
        .print-button:hover {
            background: #2980b9;
        }
        
        .collapsed {
            display: none;
        }
        
        @media print {
            body {
                background: white;
                padding: 0;
            }
            
            .container {
                box-shadow: none;
                border-radius: 0;
            }
            
            .print-button {
                display: none;
            }
            
            .section {
                page-break-inside: avoid;
            }
        }
        
        @media (max-width: 768px) {
            .report-header h1 {
                font-size: 24px;
            }
            
            .report-header h2 {
                font-size: 16px;
            }
            
            .section {
                padding: 20px;
            }
            
            .info-grid {
                grid-template-columns: 1fr;
            }
            
            .summary-grid {
                grid-template-columns: 1fr;
            }
            
            .data-table {
                font-size: 12px;
            }
            
            .data-table th,
            .data-table td {
                padding: 8px;
            }
        }
        """
        
    def _save_html_resources(self, output_dir):
        """Сохранение ресурсов для HTML отчета"""
        # Создаем папку для ресурсов
        resources_dir = os.path.join(output_dir, "resources")
        os.makedirs(resources_dir, exist_ok=True)
        
        # Сохраняем CSS в отдельный файл
        css_file = os.path.join(resources_dir, "style.css")
        with open(css_file, 'w', encoding='utf-8') as f:
            f.write(self._get_html_styles())
        
        # Сохраняем логотип (если есть)
        logo_path = os.path.join("assets", "images", "logo.png")
        if os.path.exists(logo_path):
            import shutil
            shutil.copy(logo_path, os.path.join(resources_dir, "logo.png"))


class ReportPanel(QWidget):
    """Панель управления отчетами"""
    
    report_generated = pyqtSignal(str)  # Сигнал о создании отчета
    report_error = pyqtSignal(str)      # Сигнал об ошибке
    
    def __init__(self, config_manager=None):
        super().__init__()
        self.config_manager = config_manager or ConfigManager()
        self.diagnostic_data = None
        self.report_generator = None
        self.current_report_path = None
        self.setup_ui()
        self.load_templates()
        
    def setup_ui(self):
        """Настройка интерфейса панели отчетов"""
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(10, 10, 10, 10)
        main_layout.setSpacing(10)
        
        # Заголовок панели
        title_label = QLabel("УПРАВЛЕНИЕ ОТЧЕТАМИ")
        title_label.setAlignment(Qt.AlignCenter)
        title_font = QFont()
        title_font.setPointSize(16)
        title_font.setBold(True)
        title_label.setFont(title_font)
        title_label.setStyleSheet("color: #2c3e50; padding: 10px;")
        main_layout.addWidget(title_label)
        
        # Разделитель
        separator = QFrame()
        separator.setFrameShape(QFrame.HLine)
        separator.setFrameShadow(QFrame.Sunken)
        main_layout.addWidget(separator)
        
        # Создаем табы для разных функций
        self.tab_widget = QTabWidget()
        
        # Вкладка "Быстрый отчет"
        self.quick_report_tab = self.create_quick_report_tab()
        self.tab_widget.addTab(self.quick_report_tab, "Быстрый отчет")
        
        # Вкладка "Настройки отчета"
        self.settings_tab = self.create_settings_tab()
        self.tab_widget.addTab(self.settings_tab, "Настройки")
        
        # Вкладка "История отчетов"
        self.history_tab = self.create_history_tab()
        self.tab_widget.addTab(self.history_tab, "История")
        
        # Вкладка "Шаблоны"
        self.templates_tab = self.create_templates_tab()
        self.tab_widget.addTab(self.templates_tab, "Шаблоны")
        
        main_layout.addWidget(self.tab_widget)
        
        # Статус бар
        self.status_bar = QStatusBar()
        self.status_bar.setSizeGripEnabled(False)
        main_layout.addWidget(self.status_bar)
        
        # Прогресс бар
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        main_layout.addWidget(self.progress_bar)
        
    def create_quick_report_tab(self):
        """Создание вкладки быстрого отчета"""
        tab = QWidget()
        layout = QVBoxLayout(tab)
        layout.setSpacing(15)
        
        # Информация о данных
        data_group = QGroupBox("Данные для отчета")
        data_layout = QVBoxLayout()
        
        data_info_label = QLabel(
            "Для генерации отчета необходимо выполнить диагностику.\n"
            "Текущие данные будут использованы для создания отчета."
        )
        data_info_label.setWordWrap(True)
        data_layout.addWidget(data_info_label)
        
        self.data_status_label = QLabel("Статус: Данные не загружены")
        self.data_status_label.setStyleSheet("color: #dc3545; font-weight: bold;")
        data_layout.addWidget(self.data_status_label)
        
        self.load_data_button = QPushButton("Загрузить данные диагностики")
        self.load_data_button.clicked.connect(self.load_diagnostic_data)
        data_layout.addWidget(self.load_data_button)
        
        data_group.setLayout(data_layout)
        layout.addWidget(data_group)
        
        # Настройки отчета
        report_group = QGroupBox("Настройки отчета")
        report_layout = QGridLayout()
        
        # Тип отчета
        report_layout.addWidget(QLabel("Тип отчета:"), 0, 0)
        self.report_type_combo = QComboBox()
        self.report_type_combo.addItems(["PDF", "Word (DOCX)", "Excel", "HTML"])
        report_layout.addWidget(self.report_type_combo, 0, 1)
        
        # Шаблон
        report_layout.addWidget(QLabel("Шаблон:"), 1, 0)
        self.template_combo = QComboBox()
        report_layout.addWidget(self.template_combo, 1, 1)
        
        # Выходной каталог
        report_layout.addWidget(QLabel("Сохранить в:"), 2, 0)
        self.output_path_edit = QLineEdit()
        self.output_path_edit.setText(self.config_manager.get('reports_dir', os.path.expanduser("~/Documents/NivaReports")))
        report_layout.addWidget(self.output_path_edit, 2, 1)
        
        self.browse_button = QPushButton("Обзор...")
        self.browse_button.clicked.connect(self.browse_output_dir)
        report_layout.addWidget(self.browse_button, 2, 2)
        
        report_group.setLayout(report_layout)
        layout.addWidget(report_group)
        
        # Дополнительные опции
        options_group = QGroupBox("Дополнительные опции")
        options_layout = QVBoxLayout()
        
        self.include_charts_check = QCheckBox("Включить графики параметров")
        self.include_charts_check.setChecked(True)
        options_layout.addWidget(self.include_charts_check)
        
        self.include_recommendations_check = QCheckBox("Включить рекомендации")
        self.include_recommendations_check.setChecked(True)
        options_layout.addWidget(self.include_recommendations_check)
        
        self.include_details_check = QCheckBox("Включить подробную информацию")
        self.include_details_check.setChecked(True)
        options_layout.addWidget(self.include_details_check)
        
        self.open_after_check = QCheckBox("Открыть отчет после создания")
        self.open_after_check.setChecked(True)
        options_layout.addWidget(self.open_after_check)
        
        options_group.setLayout(options_layout)
        layout.addWidget(options_group)
        
        # Кнопки управления
        button_layout = QHBoxLayout()
        
        self.generate_button = QPushButton("Сгенерировать отчет")
        self.generate_button.setIcon(QIcon("assets/icons/generate.png"))
        self.generate_button.clicked.connect(self.generate_report)
        self.generate_button.setEnabled(False)
        button_layout.addWidget(self.generate_button)
        
        self.preview_button = QPushButton("Предпросмотр")
        self.preview_button.setIcon(QIcon("assets/icons/preview.png"))
        self.preview_button.clicked.connect(self.preview_report)
        self.preview_button.setEnabled(False)
        button_layout.addWidget(self.preview_button)
        
        self.cancel_button = QPushButton("Отмена")
        self.cancel_button.setIcon(QIcon("assets/icons/cancel.png"))
        self.cancel_button.clicked.connect(self.cancel_generation)
        button_layout.addWidget(self.cancel_button)
        
        layout.addLayout(button_layout)
        
        layout.addStretch()
        return tab
        
    def create_settings_tab(self):
        """Создание вкладки настроек"""
        tab = QWidget()
        layout = QVBoxLayout(tab)
        
        # Настройки компании
        company_group = QGroupBox("Информация о компании")
        company_layout = QFormLayout()
        
        self.company_name_edit = QLineEdit()
        self.company_name_edit.setText(self.config_manager.get('company_name', ''))
        company_layout.addRow("Название компании:", self.company_name_edit)
        
        self.company_address_edit = QLineEdit()
        self.company_address_edit.setText(self.config_manager.get('company_address', ''))
        company_layout.addRow("Адрес:", self.company_address_edit)
        
        self.company_phone_edit = QLineEdit()
        self.company_phone_edit.setText(self.config_manager.get('company_phone', ''))
        company_layout.addRow("Телефон:", self.company_phone_edit)
        
        self.company_email_edit = QLineEdit()
        self.company_email_edit.setText(self.config_manager.get('company_email', ''))
        company_layout.addRow("Email:", self.company_email_edit)
        
        company_group.setLayout(company_layout)
        layout.addWidget(company_group)
        
        # Настройки отчета по умолчанию
        default_group = QGroupBox("Настройки отчетов по умолчанию")
        default_layout = QFormLayout()
        
        self.default_report_type_combo = QComboBox()
        self.default_report_type_combo.addItems(["PDF", "Word (DOCX)", "Excel", "HTML"])
        self.default_report_type_combo.setCurrentText(self.config_manager.get('default_report_type', 'PDF'))
        default_layout.addRow("Тип отчета по умолчанию:", self.default_report_type_combo)
        
        self.default_template_combo = QComboBox()
        default_layout.addRow("Шаблон по умолчанию:", self.default_template_combo)
        
        self.default_output_dir_edit = QLineEdit()
        self.default_output_dir_edit.setText(self.config_manager.get('reports_dir', os.path.expanduser("~/Documents/NivaReports")))
        default_layout.addRow("Каталог по умолчанию:", self.default_output_dir_edit)
        
        default_group.setLayout(default_layout)
        layout.addWidget(default_group)
        
        # Настройки содержимого
        content_group = QGroupBox("Содержимое отчета")
        content_layout = QVBoxLayout()
        
        self.auto_include_charts = QCheckBox("Автоматически включать графики")
        self.auto_include_charts.setChecked(self.config_manager.get('auto_include_charts', True))
        content_layout.addWidget(self.auto_include_charts)
        
        self.auto_include_recommendations = QCheckBox("Автоматически включать рекомендации")
        self.auto_include_recommendations.setChecked(self.config_manager.get('auto_include_recommendations', True))
        content_layout.addWidget(self.auto_include_recommendations)
        
        self.auto_open_report = QCheckBox("Автоматически открывать отчет после создания")
        self.auto_open_report.setChecked(self.config_manager.get('auto_open_report', True))
        content_layout.addWidget(self.auto_open_report)
        
        content_group.setLayout(content_layout)
        layout.addWidget(content_group)
        
        # Кнопки сохранения настроек
        save_layout = QHBoxLayout()
        save_button = QPushButton("Сохранить настройки")
        save_button.clicked.connect(self.save_settings)
        save_layout.addWidget(save_button)
        
        reset_button = QPushButton("Сбросить настройки")
        reset_button.clicked.connect(self.reset_settings)
        save_layout.addWidget(reset_button)
        
        layout.addLayout(save_layout)
        layout.addStretch()
        
        return tab
        
    def create_history_tab(self):
        """Создание вкладки истории отчетов"""
        tab = QWidget()
        layout = QVBoxLayout(tab)
        
        # Панель управления историей
        history_controls = QHBoxLayout()
        
        refresh_button = QPushButton("Обновить")
        refresh_button.clicked.connect(self.refresh_history)
        history_controls.addWidget(refresh_button)
        
        delete_button = QPushButton("Удалить выбранное")
        delete_button.clicked.connect(self.delete_selected_reports)
        history_controls.addWidget(delete_button)
        
        clear_all_button = QPushButton("Очистить все")
        clear_all_button.clicked.connect(self.clear_all_reports)
        clear_all_button.setStyleSheet("background-color: #dc3545; color: white;")
        history_controls.addWidget(clear_all_button)
        
        history_controls.addStretch()
        layout.addLayout(history_controls)
        
        # Таблица истории
        self.history_table = QTableWidget()
        self.history_table.setColumnCount(6)
        self.history_table.setHorizontalHeaderLabels([
            "Дата", "Тип", "Автомобиль", "Файл", "Размер", "Действия"
        ])
        self.history_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.history_table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.history_table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        
        layout.addWidget(self.history_table)
        
        # Загружаем историю
        self.load_history()
        
        return tab
        
    def create_templates_tab(self):
        """Создание вкладки управления шаблонами"""
        tab = QWidget()
        layout = QVBoxLayout(tab)
        
        # Дерево шаблонов
        self.templates_tree = QTreeWidget()
        self.templates_tree.setHeaderLabels(["Название", "Тип", "Дата изменения"])
        self.templates_tree.setColumnWidth(0, 250)
        
        layout.addWidget(self.templates_tree)
        
        # Кнопки управления шаблонами
        template_buttons = QHBoxLayout()
        
        new_template_button = QPushButton("Новый шаблон")
        new_template_button.clicked.connect(self.create_new_template)
        template_buttons.addWidget(new_template_button)
        
        edit_template_button = QPushButton("Редактировать")
        edit_template_button.clicked.connect(self.edit_template)
        template_buttons.addWidget(edit_template_button)
        
        delete_template_button = QPushButton("Удалить")
        delete_template_button.clicked.connect(self.delete_template)
        template_buttons.addWidget(delete_template_button)
        
        import_template_button = QPushButton("Импорт")
        import_template_button.clicked.connect(self.import_template)
        template_buttons.addWidget(import_template_button)
        
        export_template_button = QPushButton("Экспорт")
        export_template_button.clicked.connect(self.export_template)
        template_buttons.addWidget(export_template_button)
        
        template_buttons.addStretch()
        layout.addLayout(template_buttons)
        
        return tab
        
    def load_templates(self):
        """Загрузка шаблонов отчетов"""
        templates_dir = os.path.join("config", "templates")
        os.makedirs(templates_dir, exist_ok=True)
        
        # Стандартные шаблоны
        default_templates = [
            ("Простой отчет", "pdf"),
            ("Расширенный отчет", "pdf"),
            ("Технический отчет", "docx"),
            ("Коммерческий отчет", "docx"),
            ("Сводный отчет", "excel"),
            ("Интерактивный отчет", "html"),
        ]
        
        self.template_combo.clear()
        self.default_template_combo.clear()
        
        for name, _ in default_templates:
            self.template_combo.addItem(name)
            self.default_template_combo.addItem(name)
            
    def load_diagnostic_data(self):
        """Загрузка данных диагностики"""
        # В реальном приложении здесь будет загрузка данных из модуля диагностики
        # Для примера создаем тестовые данные
        self.diagnostic_data = self.create_sample_data()
        self.data_status_label.setText("Статус: Данные загружены")
        self.data_status_label.setStyleSheet("color: #28a745; font-weight: bold;")
        self.generate_button.setEnabled(True)
        self.preview_button.setEnabled(True)
        
    def create_sample_data(self):
        """Создание тестовых данных для демонстрации"""
        return {
            'timestamp': datetime.now().strftime('%d.%m.%Y %H:%M:%S'),
            'vehicle_model': 'Chevrolet Niva 1.7i',
            'vin': 'X9L21230012345678',
            'mileage': 125430,
            'diagnostic_status': 'Завершена',
            'systems_checked': 8,
            'errors_found': 2,
            'critical_errors': 1,
            'overall_status': 'Требует внимания',
            
            'dtcs': {
                'ENGINE': [
                    {'code': 'P0171', 'description': 'Слишком бедная смесь', 'status': 'Активная', 
                     'date': '15.12.2023', 'critical': False},
                    {'code': 'P0302', 'description': 'Пропуски зажигания в цилиндре 2', 
                     'status': 'Активная', 'date': '15.12.2023', 'critical': True},
                ],
                'ABS': [
                    {'code': 'C0121', 'description': 'Неисправность датчика скорости колеса', 
                     'status': 'Активная', 'date': '10.12.2023', 'critical': False},
                ]
            },
            
            'live_data': {
                'ENGINE_RPM': {'value': 750, 'unit': 'rpm', 'min': 650, 'max': 850, 'status': 'OK'},
                'COOLANT_TEMP': {'value': 92, 'unit': '°C', 'min': 85, 'max': 105, 'status': 'OK'},
                'INTAKE_TEMP': {'value': 35, 'unit': '°C', 'min': -40, 'max': 125, 'status': 'OK'},
                'THROTTLE_POSITION': {'value': 12.5, 'unit': '%', 'min': 0, 'max': 100, 'status': 'OK'},
                'MAF_SENSOR': {'value': 3.2, 'unit': 'g/s', 'min': 2.0, 'max': 6.0, 'status': 'OK'},
                'FUEL_PRESSURE': {'value': 350, 'unit': 'kPa', 'min': 300, 'max': 400, 'status': 'OK'},
                'O2_SENSOR_B1S1': {'value': 0.85, 'unit': 'V', 'min': 0.1, 'max': 0.9, 'status': 'Warning'},
                'BATTERY_VOLTAGE': {'value': 13.8, 'unit': 'V', 'min': 12.0, 'max': 14.5, 'status': 'OK'},
            },
            
            'recommendations': [
                {'text': 'Проверить свечи зажигания и высоковольтные провода', 
                 'priority': 'Высокий', 
                 'details': 'Пропуски зажигания могут быть вызваны изношенными свечами или поврежденными проводами.'},
                {'text': 'Проверить датчик массового расхода воздуха', 
                 'priority': 'Средний',
                 'details': 'Значение MAF близко к верхнему пределу, возможен износ датчика.'},
                {'text': 'Почистить дроссельную заслонку', 
                 'priority': 'Низкий',
                 'details': 'Положение дроссельной заслонки в норме, но профилактическая чистка рекомендуется.'},
            ],
            
            'charts': {
                'Обороты двигателя': {
                    'x_data': ['10:00', '10:01', '10:02', '10:03', '10:04'],
                    'y_data': [750, 1200, 850, 1500, 750],
                    'x_label': 'Время',
                    'y_label': 'RPM'
                },
                'Температура охлаждающей жидкости': {
                    'x_data': ['10:00', '10:01', '10:02', '10:03', '10:04'],
                    'y_data': [85, 88, 92, 90, 91],
                    'x_label': 'Время',
                    'y_label': '°C'
                }
            }
        }
        
    def browse_output_dir(self):
        """Выбор каталога для сохранения отчетов"""
        directory = QFileDialog.getExistingDirectory(
            self, 
            "Выберите каталог для сохранения отчетов",
            self.output_path_edit.text()
        )
        
        if directory:
            self.output_path_edit.setText(directory)
            
    def generate_report(self):
        """Генерация отчета"""
        if not self.diagnostic_data:
            QMessageBox.warning(self, "Ошибка", "Нет данных для генерации отчета!")
            return
            
        # Подготавливаем данные
        report_data = self.prepare_report_data()
        
        # Определяем тип отчета
        report_type_map = {
            "PDF": "pdf",
            "Word (DOCX)": "docx",
            "Excel": "excel",
            "HTML": "html"
        }
        
        report_type = report_type_map.get(self.report_type_combo.currentText(), "pdf")
        
        # Определяем расширение файла
        extensions = {
            "pdf": ".pdf",
            "docx": ".docx",
            "excel": ".xlsx",
            "html": ".html"
        }
        
        extension = extensions.get(report_type, ".pdf")
        
        # Создаем имя файла
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"diagnostic_report_{timestamp}{extension}"
        output_dir = self.output_path_edit.text()
        
        # Создаем каталог, если не существует
        os.makedirs(output_dir, exist_ok=True)
        
        output_path = os.path.join(output_dir, filename)
        
        # Начинаем генерацию в отдельном потоке
        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)
        self.generate_button.setEnabled(False)
        
        self.report_generator = ReportGeneratorThread(
            report_data, 
            report_type, 
            self.template_combo.currentText(),
            output_path
        )
        
        self.report_generator.progress.connect(self.update_progress)
        self.report_generator.status.connect(self.update_status)
        self.report_generator.finished.connect(self.on_report_finished)
        self.report_generator.error.connect(self.on_report_error)
        
        self.report_generator.start()
        
    def prepare_report_data(self):
        """Подготовка данных для отчета"""
        # Копируем данные и добавляем настройки
        report_data = self.diagnostic_data.copy()
        
        # Добавляем информацию о компании
        report_data['company'] = {
            'name': self.config_manager.get('company_name', ''),
            'address': self.config_manager.get('company_address', ''),
            'phone': self.config_manager.get('company_phone', ''),
            'email': self.config_manager.get('company_email', ''),
        }
        
        # Фильтруем данные в зависимости от настроек
        if not self.include_charts_check.isChecked():
            report_data.pop('charts', None)
            
        if not self.include_recommendations_check.isChecked():
            report_data.pop('recommendations', None)
            
        return report_data
        
    def update_progress(self, value):
        """Обновление прогресса генерации"""
        self.progress_bar.setValue(value)
        
    def update_status(self, status):
        """Обновление статуса"""
        self.status_bar.showMessage(status)
        
    def on_report_finished(self, path, success):
        """Обработка завершения генерации"""
        self.progress_bar.setVisible(False)
        self.generate_button.setEnabled(True)
        
        if success:
            self.current_report_path = path
            self.status_bar.showMessage(f"Отчет сохранен: {path}")
            
            # Добавляем в историю
            self.add_to_history(path)
            
            # Открываем отчет, если нужно
            if self.open_after_check.isChecked():
                self.open_report(path)
                
            # Показываем сообщение об успехе
            QMessageBox.information(self, "Успех", "Отчет успешно сгенерирован!")
            
            # Отправляем сигнал
            self.report_generated.emit(path)
        else:
            QMessageBox.warning(self, "Ошибка", "Не удалось сгенерировать отчет!")
            
    def on_report_error(self, error_message):
        """Обработка ошибки генерации"""
        self.progress_bar.setVisible(False)
        self.generate_button.setEnabled(True)
        
        QMessageBox.critical(self, "Ошибка", f"Ошибка при генерации отчета:\n{error_message}")
        self.report_error.emit(error_message)
        
    def cancel_generation(self):
        """Отмена генерации отчета"""
        if self.report_generator and self.report_generator.isRunning():
            reply = QMessageBox.question(
                self, 
                "Отмена генерации",
                "Вы уверены, что хотите отменить генерацию отчета?",
                QMessageBox.Yes | QMessageBox.No
            )
            
            if reply == QMessageBox.Yes:
                self.report_generator.terminate()
                self.progress_bar.setVisible(False)
                self.generate_button.setEnabled(True)
                self.status_bar.showMessage("Генерация отменена")
                
    def preview_report(self):
        """Предпросмотр отчета"""
        if not self.diagnostic_data:
            QMessageBox.warning(self, "Ошибка", "Нет данных для предпросмотра!")
            return
            
        # Создаем диалог предпросмотра
        preview_dialog = QDialog(self)
        preview_dialog.setWindowTitle("Предпросмотр отчета")
        preview_dialog.resize(900, 700)
        
        layout = QVBoxLayout(preview_dialog)
        
        # Веб-виджет для отображения HTML предпросмотра
        from PyQt5.QtWebEngineWidgets import QWebEngineView
        
        web_view = QWebEngineView()
        
        # Создаем HTML предпросмотр
        html_content = self.create_html_preview()
        web_view.setHtml(html_content)
        
        layout.addWidget(web_view)
        
        # Кнопки
        button_box = QDialogButtonBox(QDialogButtonBox.Close)
        button_box.rejected.connect(preview_dialog.reject)
        layout.addWidget(button_box)
        
        preview_dialog.exec_()
        
    def create_html_preview(self):
        """Создание HTML предпросмотра"""
        report_data = self.prepare_report_data()
        
        # Простой HTML предпросмотр
        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 20px; }}
                h1 {{ color: #2c3e50; }}
                .section {{ margin-bottom: 30px; }}
                table {{ border-collapse: collapse; width: 100%; }}
                th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                th {{ background-color: #f2f2f2; }}
                .error {{ color: red; }}
                .warning {{ color: orange; }}
                .ok {{ color: green; }}
            </style>
        </head>
        <body>
            <h1>Предпросмотр отчета</h1>
            <div class="section">
                <h2>Информация о диагностике</h2>
                <p><strong>Автомобиль:</strong> {report_data.get('vehicle_model', 'N/A')}</p>
                <p><strong>Дата:</strong> {report_data.get('timestamp', 'N/A')}</p>
                <p><strong>Пробег:</strong> {report_data.get('mileage', 'N/A')} км</p>
            </div>
            
            <div class="section">
                <h2>Ошибки ({report_data.get('errors_found', 0)})</h2>
        """
        
        if 'dtcs' in report_data and report_data['dtcs']:
            for ecu, dtc_list in report_data['dtcs'].items():
                if dtc_list:
                    html += f"<h3>{ecu}</h3>"
                    html += "<table>"
                    html += "<tr><th>Код</th><th>Описание</th><th>Статус</th></tr>"
                    for dtc in dtc_list:
                        critical_class = "error" if dtc.get('critical', False) else ""
                        html += f"<tr class='{critical_class}'>"
                        html += f"<td>{dtc.get('code', 'N/A')}</td>"
                        html += f"<td>{dtc.get('description', 'N/A')}</td>"
                        html += f"<td>{dtc.get('status', 'N/A')}</td>"
                        html += "</tr>"
                    html += "</table>"
                    
        html += """
            </div>
            
            <div class="section">
                <h2>Текущие параметры</h2>
        """
        
        if 'live_data' in report_data and report_data['live_data']:
            html += "<table>"
            html += "<tr><th>Параметр</th><th>Значение</th><th>Статус</th></tr>"
            for param_name, param_data in report_data['live_data'].items():
                status_class = param_data.get('status', 'OK').lower()
                html += f"<tr class='{status_class}'>"
                html += f"<td>{param_name}</td>"
                html += f"<td>{param_data.get('value', 'N/A')} {param_data.get('unit', '')}</td>"
                html += f"<td>{param_data.get('status', 'OK')}</td>"
                html += "</tr>"
            html += "</table>"
            
        html += """
            </div>
            
            <div class="section">
                <h2>Общее состояние: <span class="status">{}</span></h2>
            </div>
        </body>
        </html>
        """.format(report_data.get('overall_status', 'N/A'))
        
        return html
        
    def open_report(self, path):
        """Открытие сгенерированного отчета"""
        try:
            if sys.platform == "win32":
                os.startfile(path)
            elif sys.platform == "darwin":  # macOS
                os.system(f'open "{path}"')
            else:  # Linux
                os.system(f'xdg-open "{path}"')
        except Exception as e:
            QMessageBox.warning(self, "Ошибка", f"Не удалось открыть отчет: {e}")
            
    def add_to_history(self, filepath):
        """Добавление отчета в историю"""
        try:
            stats = os.stat(filepath)
            size = stats.st_size
            
            # Форматируем размер
            if size < 1024:
                size_str = f"{size} B"
            elif size < 1024 * 1024:
                size_str = f"{size/1024:.1f} KB"
            else:
                size_str = f"{size/(1024*1024):.1f} MB"
                
            # Добавляем в таблицу
            row = self.history_table.rowCount()
            self.history_table.insertRow(row)
            
            # Дата
            date_item = QTableWidgetItem(datetime.now().strftime("%d.%m.%Y %H:%M"))
            self.history_table.setItem(row, 0, date_item)
            
            # Тип
            _, ext = os.path.splitext(filepath)
            type_item = QTableWidgetItem(ext.upper().replace('.', ''))
            self.history_table.setItem(row, 1, type_item)
            
            # Автомобиль
            vehicle_item = QTableWidgetItem(self.diagnostic_data.get('vehicle_model', 'N/A'))
            self.history_table.setItem(row, 2, vehicle_item)
            
            # Файл
            filename_item = QTableWidgetItem(os.path.basename(filepath))
            self.history_table.setItem(row, 3, filename_item)
            
            # Размер
            size_item = QTableWidgetItem(size_str)
            self.history_table.setItem(row, 4, size_item)
            
            # Действия
            action_widget = QWidget()
            action_layout = QHBoxLayout(action_widget)
            action_layout.setContentsMargins(5, 2, 5, 2)
            
            open_button = QPushButton("Открыть")
            open_button.clicked.connect(lambda: self.open_report(filepath))
            open_button.setFixedSize(70, 25)
            
            show_button = QPushButton("Показать")
            show_button.clicked.connect(lambda: self.show_in_explorer(filepath))
            show_button.setFixedSize(70, 25)
            
            action_layout.addWidget(open_button)
            action_layout.addWidget(show_button)
            action_layout.addStretch()
            
            self.history_table.setCellWidget(row, 5, action_widget)
            
        except Exception as e:
            print(f"Ошибка добавления в историю: {e}")
            
    def show_in_explorer(self, filepath):
        """Показать файл в проводнике"""
        try:
            if sys.platform == "win32":
                os.system(f'explorer /select,"{filepath}"')
            elif sys.platform == "darwin":  # macOS
                os.system(f'open -R "{filepath}"')
            else:  # Linux
                os.system(f'xdg-open "{os.path.dirname(filepath)}"')
        except Exception as e:
            QMessageBox.warning(self, "Ошибка", f"Не удалось открыть проводник: {e}")
            
    def load_history(self):
        """Загрузка истории отчетов"""
        # Очищаем таблицу
        self.history_table.setRowCount(0)
        
        # В реальном приложении здесь будет загрузка из базы данных или файла
        # Для примера добавим несколько тестовых записей
        
        test_reports = [
            ("15.12.2023 14:30", "PDF", "Chevrolet Niva 1.7i", "report_20231215_1430.pdf", "2.1 MB"),
            ("10.12.2023 11:15", "DOCX", "Chevrolet Niva 1.8i", "report_20231210_1115.docx", "3.5 MB"),
            ("05.12.2023 09:45", "Excel", "Chevrolet Niva Модерн", "report_20231205_0945.xlsx", "1.8 MB"),
        ]
        
        for i, report in enumerate(test_reports):
            self.history_table.insertRow(i)
            for j, value in enumerate(report):
                item = QTableWidgetItem(value)
                self.history_table.setItem(i, j, item)
                
            # Действия для тестовых записей
            action_widget = QWidget()
            action_layout = QHBoxLayout(action_widget)
            action_layout.setContentsMargins(5, 2, 5, 2)
            
            open_button = QPushButton("Открыть")
            open_button.setFixedSize(70, 25)
            open_button.setEnabled(False)  # Тестовые файлы не существуют
            
            show_button = QPushButton("Показать")
            show_button.setFixedSize(70, 25)
            show_button.setEnabled(False)
            
            action_layout.addWidget(open_button)
            action_layout.addWidget(show_button)
            action_layout.addStretch()
            
            self.history_table.setCellWidget(i, 5, action_widget)
            
    def refresh_history(self):
        """Обновление истории отчетов"""
        self.load_history()
        self.status_bar.showMessage("История обновлена")
        
    def delete_selected_reports(self):
        """Удаление выбранных отчетов"""
        selected_rows = set()
        for item in self.history_table.selectedItems():
            selected_rows.add(item.row())
            
        if not selected_rows:
            QMessageBox.information(self, "Информация", "Не выбрано ни одного отчета для удаления")
            return
            
        reply = QMessageBox.question(
            self,
            "Подтверждение удаления",
            f"Вы уверены, что хотите удалить {len(selected_rows)} отчет(ов)?",
            QMessageBox.Yes | QMessageBox.No
        )
        
        if reply == QMessageBox.Yes:
            # Удаляем строки в обратном порядке
            for row in sorted(selected_rows, reverse=True):
                self.history_table.removeRow(row)
                
            self.status_bar.showMessage(f"Удалено {len(selected_rows)} отчет(ов)")
            
    def clear_all_reports(self):
        """Очистка всей истории отчетов"""
        if self.history_table.rowCount() == 0:
            return
            
        reply = QMessageBox.question(
            self,
            "Подтверждение очистки",
            "Вы уверены, что хотите очистить всю историю отчетов?",
            QMessageBox.Yes | QMessageBox.No
        )
        
        if reply == QMessageBox.Yes:
            self.history_table.setRowCount(0)
            self.status_bar.showMessage("История отчетов очищена")
            
    def save_settings(self):
        """Сохранение настроек"""
        settings = {
            'company_name': self.company_name_edit.text(),
            'company_address': self.company_address_edit.text(),
            'company_phone': self.company_phone_edit.text(),
            'company_email': self.company_email_edit.text(),
            'default_report_type': self.default_report_type_combo.currentText(),
            'reports_dir': self.default_output_dir_edit.text(),
            'auto_include_charts': self.auto_include_charts.isChecked(),
            'auto_include_recommendations': self.auto_include_recommendations.isChecked(),
            'auto_open_report': self.auto_open_report.isChecked(),
        }
        
        self.config_manager.update_settings(settings)
        self.config_manager.save()
        
        QMessageBox.information(self, "Успех", "Настройки сохранены")
        
    def reset_settings(self):
        """Сброс настроек к значениям по умолчанию"""
        reply = QMessageBox.question(
            self,
            "Сброс настроек",
            "Вы уверены, что хотите сбросить все настройки к значениям по умолчанию?",
            QMessageBox.Yes | QMessageBox.No
        )
        
        if reply == QMessageBox.Yes:
            self.config_manager.reset_to_defaults()
            self.load_settings()
            QMessageBox.information(self, "Успех", "Настройки сброшены")
            
    def load_settings(self):
        """Загрузка настроек"""
        # Загрузка будет выполняться при инициализации через config_manager
        pass
        
    def create_new_template(self):
        """Создание нового шаблона"""
        dialog = QDialog(self)
        dialog.setWindowTitle("Новый шаблон")
        dialog.resize(400, 300)
        
        layout = QVBoxLayout(dialog)
        
        form_layout = QFormLayout()
        
        name_edit = QLineEdit()
        form_layout.addRow("Название шаблона:", name_edit)
        
        type_combo = QComboBox()
        type_combo.addItems(["PDF", "Word", "Excel", "HTML"])
        form_layout.addRow("Тип шаблона:", type_combo)
        
        description_edit = QTextEdit()
        description_edit.setMaximumHeight(100)
        form_layout.addRow("Описание:", description_edit)
        
        layout.addLayout(form_layout)
        
        button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        button_box.accepted.connect(dialog.accept)
        button_box.rejected.connect(dialog.reject)
        layout.addWidget(button_box)
        
        if dialog.exec_() == QDialog.Accepted:
            name = name_edit.text().strip()
            if name:
                # В реальном приложении здесь будет создание файла шаблона
                QMessageBox.information(self, "Успех", f"Шаблон '{name}' создан")
                
    def edit_template(self):
        """Редактирование шаблона"""
        item = self.templates_tree.currentItem()
        if not item:
            QMessageBox.warning(self, "Ошибка", "Не выбран шаблон для редактирования")
            return
            
        template_name = item.text(0)
        QMessageBox.information(self, "Редактирование", f"Редактирование шаблона: {template_name}")
        
    def delete_template(self):
        """Удаление шаблона"""
        item = self.templates_tree.currentItem()
        if not item:
            QMessageBox.warning(self, "Ошибка", "Не выбран шаблон для удаления")
            return
            
        template_name = item.text(0)
        
        reply = QMessageBox.question(
            self,
            "Удаление шаблона",
            f"Вы уверены, что хотите удалить шаблон '{template_name}'?",
            QMessageBox.Yes | QMessageBox.No
        )
        
        if reply == QMessageBox.Yes:
            # В реальном приложении здесь будет удаление файла шаблона
            parent = item.parent()
            if parent:
                parent.removeChild(item)
            else:
                index = self.templates_tree.indexOfTopLevelItem(item)
                self.templates_tree.takeTopLevelItem(index)
                
    def import_template(self):
        """Импорт шаблона"""
        file_filter = "Шаблоны (*.json *.xml *.template);;Все файлы (*.*)"
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Импорт шаблона",
            "",
            file_filter
        )
        
        if file_path:
            try:
                # В реальном приложении здесь будет обработка импорта
                template_name = os.path.basename(file_path)
                QMessageBox.information(self, "Успех", f"Шаблон '{template_name}' импортирован")
            except Exception as e:
                QMessageBox.critical(self, "Ошибка", f"Ошибка импорта: {e}")
                
    def export_template(self):
        """Экспорт шаблона"""
        item = self.templates_tree.currentItem()
        if not item:
            QMessageBox.warning(self, "Ошибка", "Не выбран шаблон для экспорта")
            return
            
        template_name = item.text(0)
        
        file_filter = "JSON файлы (*.json);;XML файлы (*.xml);;Все файлы (*.*)"
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Экспорт шаблона",
            f"{template_name}.json",
            file_filter
        )
        
        if file_path:
            try:
                # В реальном приложении здесь будет экспорт
                QMessageBox.information(self, "Успех", f"Шаблон экспортирован в: {file_path}")
            except Exception as e:
                QMessageBox.critical(self, "Ошибка", f"Ошибка экспорта: {e}")
                
    def set_diagnostic_data(self, data):
        """Установка данных диагностики для отчета"""
        self.diagnostic_data = data
        if data:
            self.data_status_label.setText("Статус: Данные загружены")
            self.data_status_label.setStyleSheet("color: #28a745; font-weight: bold;")
            self.generate_button.setEnabled(True)
            self.preview_button.setEnabled(True)
        else:
            self.data_status_label.setText("Статус: Данные не загружены")
            self.data_status_label.setStyleSheet("color: #dc3545; font-weight: bold;")
            self.generate_button.setEnabled(False)
            self.preview_button.setEnabled(False)
            
    def get_report_types(self):
        """Возвращает список доступных типов отчетов"""
        return ["PDF", "Word (DOCX)", "Excel", "HTML"]
        
    def get_templates(self):
        """Возвращает список доступных шаблонов"""
        templates = []
        for i in range(self.template_combo.count()):
            templates.append(self.template_combo.itemText(i))
        return templates
        
    def generate_report_from_data(self, data, report_type="PDF", template="Простой отчет"):
        """Генерация отчета из предоставленных данных"""
        self.set_diagnostic_data(data)
        self.report_type_combo.setCurrentText(report_type)
        self.template_combo.setCurrentText(template)
        self.generate_report()


# Для тестирования панели
if __name__ == "__main__":
    import sys
    from PyQt5.QtWidgets import QApplication
    
    app = QApplication(sys.argv)
    window = ReportPanel()
    window.resize(1000, 700)
    window.show()
    sys.exit(app.exec_())